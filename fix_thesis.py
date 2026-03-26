#!/usr/bin/env python3
"""
Fix script for NetShield_IDS_Diploma_Thesis.docx
Fixes:
1. TOC - add page numbers with dotted leaders
2. USD -> KZT currency conversion
3. Kazakh annotation - replace transliteration with Cyrillic
4. Chapter headings - ensure page_break_before
5. Minor formatting fixes
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

DOCX_PATH = "/Users/bexul/Diplomka/NetShield_IDS_Diploma_Thesis.docx"

doc = Document(DOCX_PATH)

# ============================================================
# 1. FIX TABLE OF CONTENTS - add page numbers with dotted leaders
# ============================================================
print("=== 1. Fixing Table of Contents ===")

# Estimated page numbers mapping
toc_entries = {
    "INTRODUCTION": 8,
    "1 ANALYTICAL REVIEW": 12,
    "1.1 Overview of Network Security Threats and IDS Systems": 12,
    "1.2 Analysis of Existing IDS Approaches": 15,
    "1.3 Machine Learning Methods for Intrusion Detection": 18,
    "1.4 Review of the CICIDS2017 Dataset": 22,
    "1.5 Comparative Analysis of ML Algorithms for IDS": 25,
    "2 THEORETICAL PART": 28,
    "2.1 Autoencoder Architecture for Anomaly Detection": 28,
    "2.2 Classification Algorithms": 31,
    "2.3 Two-Level Detection Pipeline Design": 34,
    "2.4 Feature Selection and Data Preprocessing": 36,
    "2.5 Evaluation Metrics": 38,
    "3 SOFTWARE IMPLEMENTATION": 40,
    "3.1 System Architecture and Technology Stack": 40,
    "3.2 Backend Implementation": 43,
    "3.3 Frontend Dashboard": 46,
    "3.4 URL Security Analysis Module": 49,
    "3.5 Real-Time Monitoring Module": 52,
    "3.6 Model Training and Results": 54,
    "3.7 Testing and Validation": 57,
    "4 ECONOMIC EFFECTIVENESS": 60,
    "4.1 Project Cost Estimation": 60,
    "4.2 Comparison with Commercial Solutions": 62,
    "4.3 Economic Benefits": 64,
    "CONCLUSION": 66,
    "REFERENCES": 68,
    "APPENDIX A \u2013 Source Code Listings": 72,
    "APPENDIX B \u2013 Screenshots": 75,
}

# Find the TOC heading and its entries (paragraphs 100-129)
toc_heading_idx = None
toc_start = None
toc_end = None

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "TABLE OF CONTENTS":
        toc_heading_idx = i
        toc_start = i + 1
        break

if toc_heading_idx is not None:
    # Find where TOC entries end (next empty paragraph or section break before INTRODUCTION body)
    for i in range(toc_start, min(toc_start + 35, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        txt = p.text.strip()
        if txt == "" and i > toc_start + 2:
            # Check if next non-empty is not a TOC entry
            toc_end = i
            break
    if toc_end is None:
        toc_end = toc_start + 29  # fallback

    print(f"  TOC heading at index {toc_heading_idx}, entries from {toc_start} to {toc_end - 1}")

    # Process each TOC entry paragraph
    for i in range(toc_start, toc_end):
        p = doc.paragraphs[i]
        txt = p.text.strip()
        if txt in toc_entries:
            page_num = toc_entries[txt]

            # Clear existing runs
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = txt
            else:
                run = p.add_run(txt)

            # Add tab + page number
            tab_run = p.add_run()
            tab_run.font.name = "Times New Roman"
            tab_run.font.size = Pt(14)
            # Add a tab character
            tab_run.text = "\t"

            page_run = p.add_run(str(page_num))
            page_run.font.name = "Times New Roman"
            page_run.font.size = Pt(14)

            # Set up tab stop with dot leader at right margin (~16.5 cm)
            pPr = p._element.get_or_add_pPr()
            # Remove existing tabs
            for existing_tabs in pPr.findall(qn('w:tabs')):
                pPr.remove(existing_tabs)

            tabs = parse_xml(
                '<w:tabs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:tab w:val="right" w:leader="dot" w:pos="9356"/>'
                '</w:tabs>'
            )
            # 9356 twips = ~16.5 cm (right margin for A4 with standard margins)
            pPr.append(tabs)

            print(f"  Fixed TOC entry: {txt} -> page {page_num}")


# ============================================================
# 2. FIX CURRENCY: USD -> KZT
# ============================================================
print("\n=== 2. Fixing Currency (USD -> KZT) ===")


def replace_in_paragraph(para, old_text, new_text):
    """Replace text in a paragraph while preserving formatting as much as possible."""
    full_text = para.text
    if old_text not in full_text:
        return False

    # Simple approach: find which runs contain the text and replace
    # Build a mapping of character positions to runs
    runs = para.runs
    if not runs:
        return False

    # Try run-by-run replacement first
    for run in runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # If text spans multiple runs, we need a different approach
    # Concatenate all run texts, find the old_text, then redistribute
    combined = ""
    run_boundaries = []  # (start, end, run_index)
    for idx, run in enumerate(runs):
        start = len(combined)
        combined += run.text
        end = len(combined)
        run_boundaries.append((start, end, idx))

    pos = combined.find(old_text)
    if pos == -1:
        return False

    new_combined = combined[:pos] + new_text + combined[pos + len(old_text):]

    # Redistribute text across runs
    # Put all text in first run, clear others? No, better to keep formatting.
    # Simpler: clear all runs, put new text in first run
    if runs:
        runs[0].text = new_combined
        for r in runs[1:]:
            r.text = ""
    return True


def replace_in_cell(cell, old_text, new_text):
    """Replace text in a table cell."""
    for para in cell.paragraphs:
        replace_in_paragraph(para, old_text, new_text)


# --- Fix Table 13 (index 13): Development cost breakdown (Table 4.1) ---
# Headers: Phase | Duration (days) | Hourly Rate (USD) | Hours/Day | Cost (USD)
table_4_1 = doc.tables[13]
print("  Fixing Table 4.1 (Development cost breakdown)")

# Fix header row
for cell in table_4_1.rows[0].cells:
    for para in cell.paragraphs:
        if "Hourly Rate (USD)" in para.text:
            replace_in_paragraph(para, "Hourly Rate (USD)", "Hourly Rate (KZT)")
        if "Cost (USD)" in para.text:
            replace_in_paragraph(para, "Cost (USD)", "Cost (KZT)")

# Fix data rows - convert dollar values to KZT
# Hourly rate $15 -> 7,050 KZT; costs are hours * rate
cost_mapping_4_1 = {
    "15": "7,050",     # hourly rate column
    "1,350": "634,500",
    "900": "423,000",
    "450": "211,500",
    "1,800": "846,000",
    "630": "296,100",
    "720": "338,400",
    "9,900": "4,653,000",
}

for row in table_4_1.rows[1:]:
    cells = row.cells
    # Column 2: Hourly Rate
    rate_text = cells[2].text.strip()
    if rate_text == "15":
        for para in cells[2].paragraphs:
            replace_in_paragraph(para, "15", "7,050")

    # Column 4: Cost
    cost_text = cells[4].text.strip()
    if cost_text in cost_mapping_4_1:
        for para in cells[4].paragraphs:
            replace_in_paragraph(para, cost_text, cost_mapping_4_1[cost_text])

print("  Table 4.1 done")

# --- Fix Table 14 (index 14): Infrastructure cost estimation (Table 4.2) ---
# Headers: Item | Cost Type | Cost (USD/year)
table_4_2 = doc.tables[14]
print("  Fixing Table 4.2 (Infrastructure cost)")

for cell in table_4_2.rows[0].cells:
    for para in cell.paragraphs:
        if "Cost (USD/year)" in para.text:
            replace_in_paragraph(para, "Cost (USD/year)", "Cost (KZT/year)")

infra_cost_mapping = {
    "200": "94,000",
    "50-100": "23,500-47,000",
    "250-300": "117,500-141,000",
}

for row in table_4_2.rows[1:]:
    cells = row.cells
    cost_text = cells[2].text.strip()
    if cost_text in infra_cost_mapping:
        for para in cells[2].paragraphs:
            replace_in_paragraph(para, cost_text, infra_cost_mapping[cost_text])

print("  Table 4.2 done")

# --- Fix Table 15 (index 15): Cost comparison (Table 4.3) ---
# Headers: Solution | Type | Annual Cost (USD) | ML-based Detection | Real-time Dashboard
table_4_3 = doc.tables[15]
print("  Fixing Table 4.3 (Cost comparison)")

for cell in table_4_3.rows[0].cells:
    for para in cell.paragraphs:
        if "Annual Cost (USD)" in para.text:
            replace_in_paragraph(para, "Annual Cost (USD)", "Annual Cost (KZT)")

annual_cost_mapping = {
    "15,000 - 50,000": "7,050,000 - 23,500,000",
    "20,000 - 80,000": "9,400,000 - 37,600,000",
    "10,000 - 40,000": "4,700,000 - 18,800,000",
    "2,000 - 5,000": "940,000 - 2,350,000",
    "5,000 - 15,000": "2,350,000 - 7,050,000",
    "25,000 - 100,000+": "11,750,000 - 47,000,000+",
    "250 - 300": "117,500 - 141,000",
}

for row in table_4_3.rows[1:]:
    cells = row.cells
    cost_text = cells[2].text.strip()
    if cost_text in annual_cost_mapping:
        for para in cells[2].paragraphs:
            replace_in_paragraph(para, cost_text, annual_cost_mapping[cost_text])

print("  Table 4.3 done")

# --- Fix Table 16 (index 16): TCO comparison (Table 4.4) ---
# Headers: Cost Component | NetShield IDS | Snort + Custom | Cisco Firepower
table_4_4 = doc.tables[16]
print("  Fixing Table 4.4 (TCO comparison)")

tco_mapping = {
    "$10,200": "4,794,000 KZT",
    "$5,000": "2,350,000 KZT",
    "$30,000": "14,100,000 KZT",
    "$0": "0 KZT",
    "$2,000": "940,000 KZT",
    "$8,000": "3,760,000 KZT",
    "$10,000": "4,700,000 KZT",
    "$1,500": "705,000 KZT",
    "$15,000": "7,050,000 KZT",
    "$500": "235,000 KZT",
    "$3,000": "1,410,000 KZT",
    "$12,200": "5,734,000 KZT",
    "$22,000": "10,340,000 KZT",
    "$68,000": "31,960,000 KZT",
}

for row in table_4_4.rows:
    for cell in row.cells:
        cell_text = cell.text.strip()
        if cell_text in tco_mapping:
            for para in cell.paragraphs:
                replace_in_paragraph(para, cell_text, tco_mapping[cell_text])

print("  Table 4.4 done")

# --- Fix paragraphs with dollar amounts ---
print("  Fixing dollar amounts in body paragraphs...")

# Paragraph-level replacements (by content matching)
para_replacements = [
    # Paragraph 446: "$15/hour" and keep KZT reference
    ("$15/hour", "7,050 KZT/hour"),

    # Paragraph 450: "$10,200 USD (approximately 4,800,000 KZT at the exchange rate of 470 KZT/USD)"
    ("$10,200 USD (approximately 4,800,000 KZT at the exchange rate of 470 KZT/USD)",
     "4,794,000 KZT"),

    # Paragraph 455: "$2,000 to $100,000+"
    ("$2,000 to $100,000+", "940,000 KZT to 47,000,000+ KZT"),

    # Paragraph 458: "$10,000-$50,000"
    ("$10,000-$50,000", "4,700,000-23,500,000 KZT"),

    # Paragraph 460: "$3.05 million" - this is a citation from IBM report, keep as is
    # (it's a global statistic, not our project cost)

    # Paragraph 462: "estimated damage: $5,000"
    ("estimated damage: $5,000", "estimated damage: 2,350,000 KZT"),

    # Paragraph 464: ROI formula first year
    ("($5,000 - $10,200) / $10,200", "(2,350,000 - 4,794,000) / 4,794,000"),

    # Paragraph 465: ROI formula subsequent years
    ("($5,000 - $300) / $300", "(2,350,000 - 141,000) / 141,000"),

    # Paragraph 466: multiple dollar amounts
    ("$300/year", "141,000 KZT/year"),
    ("$10,200 + 4*$300 = $11,400", "4,794,000 + 4*141,000 = 5,358,000 KZT"),
    ("5*$5,000 = $25,000", "5*2,350,000 = 11,750,000 KZT"),
    ("($25,000 - $11,400) / $11,400", "(11,750,000 - 5,358,000) / 5,358,000"),

    # Paragraph 471: "$12,200"
    ("$12,200", "5,734,000 KZT"),

    # Paragraph 472: "$50-100 million"
    ("$50-100 million", "23.5-47 billion KZT"),

    # Paragraph 132: "$4.45 million" - IBM global statistic, keep as is (it's a citation)
]

for old, new in para_replacements:
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            if replace_in_paragraph(p, old, new):
                count += 1
    if count > 0:
        print(f"    Replaced '{old}' -> '{new}' in {count} paragraph(s)")
    else:
        print(f"    WARNING: '{old}' not found in any paragraph")

# Also remove remaining "USD" references in economic paragraphs (but not in abbreviations table)
# Check for leftover USD mentions
for i, p in enumerate(doc.paragraphs):
    if i < 440 or i > 500:
        continue
    if "USD" in p.text:
        replace_in_paragraph(p, "USD", "KZT")
        print(f"    Replaced remaining 'USD' in paragraph [{i}]")


# ============================================================
# 3. FIX KAZAKH ANNOTATION - replace transliteration with Cyrillic
# ============================================================
print("\n=== 3. Fixing Kazakh Annotation ===")

kazakh_heading = "АҢДАТПА"
kazakh_texts = [
    "Дипломдық жоба тақырыбы: Желілік шабуылдарды машиналық оқыту әдістерімен анықтау.",
    "Жұмыс көлемі: дипломдық жобада 4 тарау, кіріспе, қорытынды, әдебиеттер тізімі және қосымшалар бар.",
    "Негізгі сөз тіркестері: кіру анықтау жүйесі (IDS), автоэнкодер, аномалияны анықтау, XGBoost, LightGBM, Random Forest, нейрондық желі, CICIDS2017, желілік қауіпсіздік, нақты уақыттағы мониторинг, URL қауіпсіздігін талдау, OSINT, VirusTotal.",
    "",  # empty paragraph (index 74)
    "Дипломдық жобаның мақсаты — желілік трафикті талдау арқылы кибершабуылдарды автоматты анықтау үшін екі деңгейлі кіру анықтау жүйесін (IDS) жасау. Бірінші деңгейде PyTorch негізінде жасалған автоэнкодер нейрондық желісі қалыпты трафиктің реконструкция қателігін (reconstruction error) пайдаланып аномалияларды анықтайды. Екінші деңгейде XGBoost, LightGBM, Random Forest және MLP классификаторларының ансамблі шабуыл түрін нақтылайды.",
    "",  # empty paragraph (index 76)
    "Жоба сонымен қатар URL қауіпсіздігін талдау модулін (OSINT threat intelligence фидтерімен және VirusTotal API интеграциясымен) және WebSocket арқылы нақты уақыттағы мониторингті жүзеге асырады. Жүйе FastAPI бэкендін, React фронтенд панелін және SQLite дерекқорын қамтиды. CICIDS2017 деректер жиынтығында тестілеу кезінде жүйе 99.87% accuracy және 99.86% F1-score көрсеткіштеріне қол жеткізді.",
]

# Find ANDATPA heading
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ANDATPA":
        print(f"  Found ANDATPA heading at index {i}")

        # Replace heading text
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = kazakh_heading
        else:
            run = p.add_run(kazakh_heading)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        print(f"  Replaced heading: ANDATPA -> {kazakh_heading}")

        # Replace body paragraphs (indices 71-77)
        for j, new_text in enumerate(kazakh_texts):
            para_idx = i + 1 + j
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                elif new_text:
                    run = para.add_run(new_text)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)
                if new_text:
                    print(f"  Replaced paragraph [{para_idx}]: {new_text[:60]}...")
        break


# ============================================================
# 4. FIX CHAPTER HEADINGS - page_break_before
# ============================================================
print("\n=== 4. Fixing Chapter Headings (page_break_before) ===")

# The chapter headings that should start on a new page.
# From the analysis, page breaks already exist in preceding empty paragraphs.
# We'll also set page_break_before on the heading paragraph itself for robustness.
chapter_heading_texts = [
    "INTRODUCTION",
    "1 ANALYTICAL REVIEW",
    "2 THEORETICAL PART",
    "3 SOFTWARE IMPLEMENTATION",
    "4 ECONOMIC EFFECTIVENESS",
    "CONCLUSION",
    "REFERENCES",
]

# We need to match the actual content headings (not TOC entries)
# The TOC entries are at indices 101-129, actual headings are at 131+
# We'll match paragraphs that are centered, bold, and match the text

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    # Skip TOC entries (indices 100-130 approximately)
    if 100 <= i <= 130:
        continue
    if txt in chapter_heading_texts:
        # Check if it looks like a heading (bold, centered)
        is_bold = p.runs and p.runs[0].bold
        is_centered = p.alignment == WD_ALIGN_PARAGRAPH.CENTER
        if is_bold or is_centered:
            p.paragraph_format.page_break_before = True
            print(f"  Set page_break_before for [{i}] '{txt}'")

# Also handle APPENDIX headings
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt.startswith("APPENDIX") and i > 130:
        is_bold = p.runs and p.runs[0].bold
        if is_bold:
            p.paragraph_format.page_break_before = True
            print(f"  Set page_break_before for [{i}] '{txt}'")


# ============================================================
# 5. FIX MINOR ISSUES
# ============================================================
print("\n=== 5. Fixing Minor Issues ===")

# 5a. TABLE OF CONTENTS heading - already bold, centered, 14pt. Verify.
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "TABLE OF CONTENTS":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
        print(f"  Ensured TABLE OF CONTENTS heading formatting at [{i}]")
        break

# 5b. АННОТАЦИЯ and ANNOTATION headings - ensure centered and bold
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt in ["АННОТАЦИЯ", "ANNOTATION", "АҢДАТПА"]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
        print(f"  Ensured formatting for [{i}] '{txt}'")


# ============================================================
# SAVE
# ============================================================
print("\n=== Saving document ===")
doc.save(DOCX_PATH)
print(f"Saved to {DOCX_PATH}")
print("Done!")
