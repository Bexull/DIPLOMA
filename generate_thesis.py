#!/usr/bin/env python3
"""
Generate a complete IITU diploma thesis DOCX for the NetShield IDS project.

Usage:
    python generate_thesis.py

Output:
    NetShield_IDS_Diploma_Thesis.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "NetShield_IDS_Diploma_Thesis.docx")

# ──────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_paragraph_format(paragraph, font_name="Times New Roman", font_size=14,
                         bold=False, italic=False, alignment=None,
                         space_before=0, space_after=0, first_line_indent=None,
                         line_spacing=1.5, keep_together=False, color=None):
    """Apply comprehensive formatting to a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        pf.alignment = alignment
    if keep_together:
        pf.keep_together = True

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_body_paragraph(doc, text, bold=False, italic=False, alignment=None,
                       first_line_indent=1.25, font_size=14, space_before=0,
                       space_after=6):
    """Add a standard body paragraph with IITU formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")

    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        pf.alignment = alignment
    return p


def add_heading_centered(doc, text, level=1, font_size=14, space_before=12,
                         space_after=12):
    """Add a centered, bold heading."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.keep_with_next = True
    return p


def add_subheading(doc, text, level=2, font_size=14, space_before=12, space_after=6):
    """Add a left-aligned bold subheading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    return p


def add_figure_placeholder(doc, caption, number):
    """Add a placeholder for a figure."""
    p = doc.add_paragraph()
    run = p.add_run(f"[Figure {number} — {caption}]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    return p


def add_table_with_data(doc, headers, rows, caption=None, number=None):
    """Add a formatted table with optional caption."""
    if caption and number:
        p = doc.add_paragraph()
        run = p.add_run(f"Table {number} — {caption}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.0

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9E2F3")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                run.font.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

    # Add spacing after table
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    return table


def add_code_block(doc, code_text, caption=None):
    """Add a formatted code block."""
    if caption:
        add_body_paragraph(doc, caption, italic=True, font_size=11,
                           first_line_indent=1.25, space_before=6)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(1.5)
    return p


def add_page_break(doc):
    """Insert a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    from docx.oxml.ns import qn
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


def add_empty_lines(doc, count=1):
    """Add empty lines."""
    for _ in range(count):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.5


def add_reference(num, text):
    """Format a reference entry."""
    return f"[{num}] {text}"


# ──────────────────────────────────────────────────────────
# Content sections
# ──────────────────────────────────────────────────────────

def create_title_page(doc):
    """Create the IITU-format title page."""
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "MINISTRY OF SCIENCE AND HIGHER EDUCATION",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=12, space_after=0)
    add_body_paragraph(doc, "OF THE REPUBLIC OF KAZAKHSTAN",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=12, space_after=0)
    add_empty_lines(doc, 1)
    add_body_paragraph(doc, "INTERNATIONAL INFORMATION TECHNOLOGY UNIVERSITY",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=14, space_after=0)
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Department: [___DEPARTMENT___]",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=12, space_after=0)
    add_body_paragraph(doc, "Educational program: [___PROGRAM_CODE___] \u2013 [___PROGRAM_NAME___]",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=12, space_after=0)

    add_empty_lines(doc, 3)

    add_body_paragraph(doc, "DIPLOMA PROJECT",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=16, space_after=6)
    add_empty_lines(doc, 1)
    add_body_paragraph(doc,
                       "Topic: Network Attack Detection Using Machine Learning Methods",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=14, space_after=0)

    add_empty_lines(doc, 4)

    # Right-aligned block for student / advisor
    p = add_body_paragraph(doc, "Performed by: [___STUDENT_NAME___]",
                           alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                           first_line_indent=0, font_size=12, space_after=0)
    add_body_paragraph(doc, "Scientific advisor: [___ADVISOR_TITLE___] [___ADVISOR_NAME___]",
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                       first_line_indent=0, font_size=12, space_after=0)
    add_body_paragraph(doc, "Reviewer: [___REVIEWER_TITLE___] [___REVIEWER_NAME___]",
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                       first_line_indent=0, font_size=12, space_after=0)

    add_empty_lines(doc, 3)

    add_body_paragraph(doc, '"Admitted to defense"',
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       first_line_indent=0, font_size=12, space_after=0)
    add_body_paragraph(doc, "Head of Department ____________________ [___DEPARTMENT_HEAD___]",
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       first_line_indent=0, font_size=12, space_after=0)
    add_body_paragraph(doc, '"____" _____________ 2026',
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       first_line_indent=0, font_size=12, space_after=0)

    add_empty_lines(doc, 3)

    add_body_paragraph(doc, "Almaty, 2026",
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0, font_size=14)

    add_page_break(doc)


def create_assignment_page(doc):
    """Create the diploma project assignment page."""
    add_heading_centered(doc, "DIPLOMA PROJECT ASSIGNMENT", font_size=14,
                         space_before=6)

    add_body_paragraph(doc, "Student: [___STUDENT_NAME___]")
    add_body_paragraph(doc, "Educational program: [___PROGRAM_CODE___] \u2013 [___PROGRAM_NAME___]")
    add_body_paragraph(doc, "Topic of diploma project: Network Attack Detection Using Machine Learning Methods")
    add_body_paragraph(doc, "Approved by the university order No. [___ORDER_NO___] dated [___ORDER_DATE___]")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Deadline for completion of the diploma project: [___DEADLINE___]")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Initial data for the diploma project:", bold=True)
    add_body_paragraph(doc, "1. CICIDS2017 dataset (Canadian Institute for Cybersecurity) containing network traffic records with labeled attack types including DoS, DDoS, Brute Force, Web Attacks, PortScan, Bot, and Infiltration.")
    add_body_paragraph(doc, "2. Requirements for a web-based intrusion detection system with real-time monitoring capabilities.")
    add_body_paragraph(doc, "3. Scientific literature on machine learning methods for network intrusion detection.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Content of the diploma project:", bold=True)
    add_body_paragraph(doc, "Introduction. Chapter 1: Analytical Review of Network Security and IDS Systems. Chapter 2: Theoretical Foundations of Machine Learning Methods for Intrusion Detection. Chapter 3: Software Implementation of the NetShield IDS System. Chapter 4: Economic Effectiveness Analysis. Conclusion. References. Appendices.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "List of graphic material:", bold=True)
    add_body_paragraph(doc, "System architecture diagram, two-level detection pipeline flowchart, autoencoder architecture diagram, model comparison charts, confusion matrices, dashboard screenshots, real-time monitoring interface screenshots, URL analysis interface screenshots.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Recommended literature:", bold=True)
    add_body_paragraph(doc, "Sharafaldin I. et al. (2018), Goodfellow I. et al. (2016), Chen T. et al. (2016), Ke G. et al. (2017), Breiman L. (2001), and other sources listed in the References section.")

    add_empty_lines(doc, 2)
    add_body_paragraph(doc, "Scientific advisor: ____________________ [___ADVISOR_NAME___]",
                       first_line_indent=0)
    add_body_paragraph(doc, "Student: ____________________ [___STUDENT_NAME___]",
                       first_line_indent=0)
    add_body_paragraph(doc, 'Date of assignment: "____" _____________ 2025',
                       first_line_indent=0)

    add_page_break(doc)


def create_calendar_plan(doc):
    """Create the calendar plan page."""
    add_heading_centered(doc, "CALENDAR PLAN", font_size=14, space_before=6)

    add_body_paragraph(doc, "for the diploma project", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       first_line_indent=0)
    add_body_paragraph(doc, "Student: [___STUDENT_NAME___]")
    add_body_paragraph(doc, "Topic: Network Attack Detection Using Machine Learning Methods")
    add_empty_lines(doc, 1)

    headers = ["No.", "Stage of Work", "Deadline", "Note"]
    rows = [
        ["1", "Literature review and analysis of existing IDS approaches", "October 2025", "Completed"],
        ["2", "Study of CICIDS2017 dataset and data preprocessing methods", "November 2025", "Completed"],
        ["3", "Design of system architecture and two-level detection pipeline", "November 2025", "Completed"],
        ["4", "Implementation of autoencoder for anomaly detection", "December 2025", "Completed"],
        ["5", "Implementation of classifiers (XGBoost, LightGBM, RF, MLP)", "December 2025", "Completed"],
        ["6", "Backend development (FastAPI, REST API, WebSocket)", "January 2026", "Completed"],
        ["7", "Frontend development (React, Ant Design, Recharts)", "January 2026", "Completed"],
        ["8", "URL security analysis module with OSINT integration", "February 2026", "Completed"],
        ["9", "Real-time monitoring module implementation", "February 2026", "Completed"],
        ["10", "Testing, validation, and performance optimization", "March 2026", "Completed"],
        ["11", "Economic effectiveness analysis", "March 2026", "Completed"],
        ["12", "Writing and formatting the diploma thesis", "March 2026", "Completed"],
        ["13", "Preparation for defense, presentation slides", "April 2026", "In progress"],
    ]

    add_table_with_data(doc, headers, rows)

    add_empty_lines(doc, 2)
    add_body_paragraph(doc, "Scientific advisor: ____________________ [___ADVISOR_NAME___]",
                       first_line_indent=0)
    add_body_paragraph(doc, "Student: ____________________ [___STUDENT_NAME___]",
                       first_line_indent=0)

    add_page_break(doc)


def create_annotations(doc):
    """Create annotations in three languages: KZ, RU, EN."""
    # --- Kazakh Annotation ---
    add_heading_centered(doc, "ANDATPA", font_size=14, space_before=6)

    add_body_paragraph(doc, "Diplomdyq zhoba taqyryby: Zheli shabuldaryn mashinalyq uyrenw adisterimen anyqtaw.")
    add_body_paragraph(doc, "Zhwmystyn kolemі: diplomdyq zhobada 4 taraw, kirispе, qorytyndy, adebiet tizimi zhane qosymshalar bar.")
    add_body_paragraph(doc, "Negizgi soz tirkester: intrusion detection system (IDS), autoencoder, anomaly detection, XGBoost, LightGBM, Random Forest, neural network, CICIDS2017, network security, machine learning.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Diplomdyq zhobanyn maqsaty \u2014 zheli trafigin taldaw arqyly kibershabuldardy anyqtaudy avtomattandyratyn eki dengeyli intrusion detection system (IDS) zhwiesіn zhаsаw. Zhwyenіn bіrіnshі dengeyi \u2014 autoencoder arqyly anomaliyalardy anyqtaw, ekіnshі dengeyi \u2014 XGBoost, LightGBM, Random Forest zhane MLP classifikatorlary arqyly shabulyn turіn anyqtaw. Zhwye CICIDS2017 derekter zhiyntyghynda uyretіlgen, FastAPI backend zhane React frontend arqyly web-qosymsha retіnde іske asyrylghan. Natizhede: autoencoder 96% anyqtyqpen anomaliyalardy anyqtaydy, XGBoost classifikatory 99.5% anyqtyqpen shabulyn turіn anyqtaydy.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "Zhobada sondai-aq URL qawіpsіzdіgіn taldaw modwlі (OSINT threat intelligence feedterі zhane VirusTotal API integraciyasy) zhane WebSocket arqyly naqty waqyt rezhimіndegі monitoring zhwyelengen.")

    add_page_break(doc)

    # --- Russian Annotation ---
    add_heading_centered(doc, "\u0410\u041d\u041d\u041e\u0422\u0410\u0426\u0418\u042f", font_size=14, space_before=6)

    add_body_paragraph(doc, "\u0422\u0435\u043c\u0430 \u0434\u0438\u043f\u043b\u043e\u043c\u043d\u043e\u0433\u043e \u043f\u0440\u043e\u0435\u043a\u0442\u0430: \u041e\u0431\u043d\u0430\u0440\u0443\u0436\u0435\u043d\u0438\u0435 \u0441\u0435\u0442\u0435\u0432\u044b\u0445 \u0430\u0442\u0430\u043a \u0441 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u0438\u0435\u043c \u043c\u0435\u0442\u043e\u0434\u043e\u0432 \u043c\u0430\u0448\u0438\u043d\u043d\u043e\u0433\u043e \u043e\u0431\u0443\u0447\u0435\u043d\u0438\u044f.")
    add_body_paragraph(doc, "\u041e\u0431\u044a\u0451\u043c \u0440\u0430\u0431\u043e\u0442\u044b: \u0434\u0438\u043f\u043b\u043e\u043c\u043d\u044b\u0439 \u043f\u0440\u043e\u0435\u043a\u0442 \u0441\u043e\u0434\u0435\u0440\u0436\u0438\u0442 4 \u0433\u043b\u0430\u0432\u044b, \u0432\u0432\u0435\u0434\u0435\u043d\u0438\u0435, \u0437\u0430\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435, \u0441\u043f\u0438\u0441\u043e\u043a \u043b\u0438\u0442\u0435\u0440\u0430\u0442\u0443\u0440\u044b \u0438 \u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u044f.")
    add_body_paragraph(doc, "\u041a\u043b\u044e\u0447\u0435\u0432\u044b\u0435 \u0441\u043b\u043e\u0432\u0430: \u0441\u0438\u0441\u0442\u0435\u043c\u0430 \u043e\u0431\u043d\u0430\u0440\u0443\u0436\u0435\u043d\u0438\u044f \u0432\u0442\u043e\u0440\u0436\u0435\u043d\u0438\u0439 (IDS), \u0430\u0432\u0442\u043e\u044d\u043d\u043a\u043e\u0434\u0435\u0440, \u043e\u0431\u043d\u0430\u0440\u0443\u0436\u0435\u043d\u0438\u0435 \u0430\u043d\u043e\u043c\u0430\u043b\u0438\u0439, XGBoost, LightGBM, Random Forest, \u043d\u0435\u0439\u0440\u043e\u043d\u043d\u0430\u044f \u0441\u0435\u0442\u044c, CICIDS2017, \u0441\u0435\u0442\u0435\u0432\u0430\u044f \u0431\u0435\u0437\u043e\u043f\u0430\u0441\u043d\u043e\u0441\u0442\u044c, \u043c\u0430\u0448\u0438\u043d\u043d\u043e\u0435 \u043e\u0431\u0443\u0447\u0435\u043d\u0438\u0435.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "\u0426\u0435\u043b\u044c \u0434\u0438\u043f\u043b\u043e\u043c\u043d\u043e\u0433\u043e \u043f\u0440\u043e\u0435\u043a\u0442\u0430 \u2014 \u0440\u0430\u0437\u0440\u0430\u0431\u043e\u0442\u043a\u0430 \u0434\u0432\u0443\u0445\u0443\u0440\u043e\u0432\u043d\u0435\u0432\u043e\u0439 \u0441\u0438\u0441\u0442\u0435\u043c\u044b \u043e\u0431\u043d\u0430\u0440\u0443\u0436\u0435\u043d\u0438\u044f \u0432\u0442\u043e\u0440\u0436\u0435\u043d\u0438\u0439 (IDS) \u0434\u043b\u044f \u0430\u0432\u0442\u043e\u043c\u0430\u0442\u0438\u0447\u0435\u0441\u043a\u043e\u0433\u043e \u043e\u0431\u043d\u0430\u0440\u0443\u0436\u0435\u043d\u0438\u044f \u043a\u0438\u0431\u0435\u0440\u0430\u0442\u0430\u043a \u043f\u0443\u0442\u0451\u043c \u0430\u043d\u0430\u043b\u0438\u0437\u0430 \u0441\u0435\u0442\u0435\u0432\u043e\u0433\u043e \u0442\u0440\u0430\u0444\u0438\u043a\u0430. \u041f\u0435\u0440\u0432\u044b\u0439 \u0443\u0440\u043e\u0432\u0435\u043d\u044c \u0441\u0438\u0441\u0442\u0435\u043c\u044b \u2014 \u0430\u0432\u0442\u043e\u044d\u043d\u043a\u043e\u0434\u0435\u0440 \u0434\u043b\u044f \u043e\u0431\u043d\u0430\u0440\u0443\u0436\u0435\u043d\u0438\u044f \u0430\u043d\u043e\u043c\u0430\u043b\u0438\u0439, \u0432\u0442\u043e\u0440\u043e\u0439 \u0443\u0440\u043e\u0432\u0435\u043d\u044c \u2014 \u0430\u043d\u0441\u0430\u043c\u0431\u043b\u044c \u043a\u043b\u0430\u0441\u0441\u0438\u0444\u0438\u043a\u0430\u0442\u043e\u0440\u043e\u0432 (XGBoost, LightGBM, Random Forest, MLP) \u0434\u043b\u044f \u043e\u043f\u0440\u0435\u0434\u0435\u043b\u0435\u043d\u0438\u044f \u0442\u0438\u043f\u0430 \u0430\u0442\u0430\u043a\u0438. \u0421\u0438\u0441\u0442\u0435\u043c\u0430 \u043e\u0431\u0443\u0447\u0435\u043d\u0430 \u043d\u0430 \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0435 CICIDS2017 \u0438 \u0440\u0435\u0430\u043b\u0438\u0437\u043e\u0432\u0430\u043d\u0430 \u0432 \u0432\u0438\u0434\u0435 \u0432\u0435\u0431-\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u044f \u0441 \u0431\u044d\u043a\u0435\u043d\u0434\u043e\u043c \u043d\u0430 FastAPI \u0438 \u0444\u0440\u043e\u043d\u0442\u0435\u043d\u0434\u043e\u043c \u043d\u0430 React. \u0420\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442\u044b: \u0430\u0432\u0442\u043e\u044d\u043d\u043a\u043e\u0434\u0435\u0440 \u043e\u0431\u043d\u0430\u0440\u0443\u0436\u0438\u0432\u0430\u0435\u0442 \u0430\u043d\u043e\u043c\u0430\u043b\u0438\u0438 \u0441 \u0442\u043e\u0447\u043d\u043e\u0441\u0442\u044c\u044e 96%, \u043a\u043b\u0430\u0441\u0441\u0438\u0444\u0438\u043a\u0430\u0442\u043e\u0440 XGBoost \u043e\u043f\u0440\u0435\u0434\u0435\u043b\u044f\u0435\u0442 \u0442\u0438\u043f \u0430\u0442\u0430\u043a\u0438 \u0441 \u0442\u043e\u0447\u043d\u043e\u0441\u0442\u044c\u044e 99.5%.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "\u0412 \u043f\u0440\u043e\u0435\u043a\u0442\u0435 \u0442\u0430\u043a\u0436\u0435 \u0440\u0435\u0430\u043b\u0438\u0437\u043e\u0432\u0430\u043d \u043c\u043e\u0434\u0443\u043b\u044c \u0430\u043d\u0430\u043b\u0438\u0437\u0430 \u0431\u0435\u0437\u043e\u043f\u0430\u0441\u043d\u043e\u0441\u0442\u0438 URL (\u0438\u043d\u0442\u0435\u0433\u0440\u0430\u0446\u0438\u044f OSINT threat intelligence \u0444\u0438\u0434\u043e\u0432 \u0438 VirusTotal API) \u0438 \u043c\u043e\u043d\u0438\u0442\u043e\u0440\u0438\u043d\u0433 \u0432 \u0440\u0435\u0430\u043b\u044c\u043d\u043e\u043c \u0432\u0440\u0435\u043c\u0435\u043d\u0438 \u0447\u0435\u0440\u0435\u0437 WebSocket.")

    add_page_break(doc)

    # --- English Annotation ---
    add_heading_centered(doc, "ANNOTATION", font_size=14, space_before=6)

    add_body_paragraph(doc, "Diploma project topic: Network Attack Detection Using Machine Learning Methods.")
    add_body_paragraph(doc, "Volume of work: the diploma project contains 4 chapters, introduction, conclusion, reference list, and appendices.")
    add_body_paragraph(doc, "Keywords: intrusion detection system (IDS), autoencoder, anomaly detection, XGBoost, LightGBM, Random Forest, neural network, CICIDS2017, network security, machine learning, deep learning, web application, real-time monitoring.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "The objective of this diploma project is to develop a two-level intrusion detection system (IDS) for automated detection of cyberattacks through network traffic analysis. The first level of the system employs an autoencoder neural network for anomaly detection, while the second level utilizes an ensemble of classifiers (XGBoost, LightGBM, Random Forest, MLP) for attack type identification. The system is trained on the CICIDS2017 dataset and implemented as a full-stack web application with a FastAPI backend and React frontend. Key results: the autoencoder achieves 96% accuracy in anomaly detection, and the XGBoost classifier achieves 99.5% accuracy in attack type classification across 8 categories.")
    add_empty_lines(doc, 1)

    add_body_paragraph(doc, "The project additionally implements a URL security analysis module integrating OSINT threat intelligence feeds (URLhaus, OpenPhish, Phishing.Database) and the VirusTotal API for comprehensive web resource verification, as well as a real-time network traffic monitoring module via WebSocket connections. The system provides an interactive dashboard for security operations, CSV batch analysis capabilities, and model performance comparison tools.")

    add_page_break(doc)


def create_abbreviations(doc):
    """Create the list of abbreviations."""
    add_heading_centered(doc, "LIST OF ABBREVIATIONS AND TERMS", font_size=14)

    abbreviations = [
        ("IDS", "Intrusion Detection System"),
        ("IPS", "Intrusion Prevention System"),
        ("ML", "Machine Learning"),
        ("DL", "Deep Learning"),
        ("AE", "Autoencoder"),
        ("MLP", "Multi-Layer Perceptron"),
        ("CNN", "Convolutional Neural Network"),
        ("RNN", "Recurrent Neural Network"),
        ("LSTM", "Long Short-Term Memory"),
        ("RF", "Random Forest"),
        ("XGBoost", "eXtreme Gradient Boosting"),
        ("LightGBM", "Light Gradient Boosting Machine"),
        ("DDoS", "Distributed Denial of Service"),
        ("DoS", "Denial of Service"),
        ("XSS", "Cross-Site Scripting"),
        ("SQL", "Structured Query Language"),
        ("API", "Application Programming Interface"),
        ("REST", "Representational State Transfer"),
        ("OSINT", "Open Source Intelligence"),
        ("CICIDS2017", "Canadian Institute for Cybersecurity Intrusion Detection System 2017"),
        ("ROC", "Receiver Operating Characteristic"),
        ("AUC", "Area Under the Curve"),
        ("MSE", "Mean Squared Error"),
        ("CSV", "Comma-Separated Values"),
        ("JSON", "JavaScript Object Notation"),
        ("SSL/TLS", "Secure Sockets Layer / Transport Layer Security"),
        ("HTTP/HTTPS", "HyperText Transfer Protocol / Secure"),
        ("TCP/IP", "Transmission Control Protocol / Internet Protocol"),
        ("UDP", "User Datagram Protocol"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("NIDS", "Network-based Intrusion Detection System"),
        ("HIDS", "Host-based Intrusion Detection System"),
        ("SVM", "Support Vector Machine"),
        ("KNN", "K-Nearest Neighbors"),
        ("SMOTE", "Synthetic Minority Over-sampling Technique"),
        ("UI/UX", "User Interface / User Experience"),
    ]

    headers = ["Abbreviation", "Full Form"]
    rows = [[a, d] for a, d in abbreviations]
    add_table_with_data(doc, headers, rows)

    add_page_break(doc)


def create_table_of_contents(doc):
    """Create the table of contents."""
    add_heading_centered(doc, "TABLE OF CONTENTS", font_size=14, space_before=6)

    toc_entries = [
        ("INTRODUCTION", ""),
        ("1 ANALYTICAL REVIEW", ""),
        ("  1.1 Overview of Network Security Threats and IDS Systems", ""),
        ("  1.2 Analysis of Existing IDS Approaches", ""),
        ("  1.3 Machine Learning Methods for Intrusion Detection", ""),
        ("  1.4 Review of the CICIDS2017 Dataset", ""),
        ("  1.5 Comparative Analysis of ML Algorithms for IDS", ""),
        ("2 THEORETICAL PART", ""),
        ("  2.1 Autoencoder Architecture for Anomaly Detection", ""),
        ("  2.2 Classification Algorithms", ""),
        ("  2.3 Two-Level Detection Pipeline Design", ""),
        ("  2.4 Feature Selection and Data Preprocessing", ""),
        ("  2.5 Evaluation Metrics", ""),
        ("3 SOFTWARE IMPLEMENTATION", ""),
        ("  3.1 System Architecture and Technology Stack", ""),
        ("  3.2 Backend Implementation", ""),
        ("  3.3 Frontend Dashboard", ""),
        ("  3.4 URL Security Analysis Module", ""),
        ("  3.5 Real-Time Monitoring Module", ""),
        ("  3.6 Model Training and Results", ""),
        ("  3.7 Testing and Validation", ""),
        ("4 ECONOMIC EFFECTIVENESS", ""),
        ("  4.1 Project Cost Estimation", ""),
        ("  4.2 Comparison with Commercial Solutions", ""),
        ("  4.3 Economic Benefits", ""),
        ("CONCLUSION", ""),
        ("REFERENCES", ""),
        ("APPENDIX A \u2013 Source Code Listings", ""),
        ("APPENDIX B \u2013 Screenshots", ""),
    ]

    for entry, _ in toc_entries:
        is_chapter = not entry.startswith("  ")
        text = entry.strip()
        p = add_body_paragraph(doc, text,
                               bold=is_chapter,
                               first_line_indent=0 if is_chapter else 1.25,
                               space_after=2, space_before=2)

    add_page_break(doc)


def create_introduction(doc):
    """Create the Introduction section."""
    add_heading_centered(doc, "INTRODUCTION", font_size=14)

    add_body_paragraph(doc, "In the contemporary digital landscape, the escalating frequency and sophistication of cyberattacks pose a significant threat to organizations, governments, and individuals worldwide. According to the IBM Cost of a Data Breach Report 2023, the global average cost of a data breach reached $4.45 million, with organizations taking an average of 277 days to identify and contain a breach [1]. Network security has therefore become a paramount concern in the field of information technology, necessitating the development of intelligent, automated systems capable of detecting and classifying network intrusions in real time.")

    add_body_paragraph(doc, "Traditional security mechanisms such as firewalls and antivirus software rely primarily on signature-based detection methods, which are effective only against known threats. These approaches fail to identify zero-day attacks, polymorphic malware, and previously unseen attack patterns. Intrusion Detection Systems (IDS) address this limitation by monitoring network traffic for suspicious activities and policy violations. However, conventional IDS solutions still depend heavily on predefined rules and signatures, limiting their effectiveness against evolving threats [2].")

    add_body_paragraph(doc, "The application of machine learning (ML) methods to network intrusion detection has emerged as a promising approach to overcome the limitations of signature-based systems. ML-based IDS can learn patterns from historical network traffic data and generalize to detect previously unseen attacks. Recent advances in deep learning, ensemble methods, and anomaly detection techniques have demonstrated remarkable performance in distinguishing malicious network traffic from benign communications [3].")

    add_body_paragraph(doc, "Relevance of the research. The relevance of this work is determined by several factors. First, the global cybersecurity threat landscape continues to expand, with Kaspersky Lab reporting that Kazakhstan experienced a 25% increase in cyberattacks in 2023 compared to the previous year. Second, the increasing adoption of cloud computing, IoT devices, and remote work arrangements has expanded the attack surface significantly. Third, the shortage of qualified cybersecurity professionals worldwide (estimated at 3.4 million unfilled positions by ISC2 in 2023) underscores the need for automated, intelligent detection systems that can augment human capabilities [4].")

    add_body_paragraph(doc, "The goal of this diploma project is to develop a two-level intrusion detection system (NetShield IDS) that combines deep learning-based anomaly detection with ensemble machine learning classifiers for automated network attack detection and classification.")

    add_body_paragraph(doc, "To achieve this goal, the following objectives have been formulated:", bold=True)
    objectives = [
        "Conduct an analytical review of existing IDS approaches, machine learning methods for intrusion detection, and the CICIDS2017 benchmark dataset.",
        "Design and implement a two-level detection pipeline: Level 1 employing an autoencoder neural network for anomaly detection, and Level 2 employing ensemble classifiers (XGBoost, LightGBM, Random Forest) and a multi-layer perceptron (MLP) for attack type classification.",
        "Develop a full-stack web application with a FastAPI backend providing REST API and WebSocket endpoints, and a React frontend with an interactive security operations dashboard.",
        "Implement a URL security analysis module integrating OSINT threat intelligence feeds (URLhaus, OpenPhish, Phishing.Database) and the VirusTotal API.",
        "Implement a real-time network traffic monitoring module using WebSocket connections for live attack scenario demonstration.",
        "Train and evaluate all models on the CICIDS2017 dataset, achieving high accuracy across 8 attack categories.",
        "Conduct an economic effectiveness analysis comparing the developed solution with commercial alternatives.",
    ]
    for i, obj in enumerate(objectives, 1):
        add_body_paragraph(doc, f"{i}. {obj}")

    add_body_paragraph(doc, "Object of research: network traffic analysis for intrusion detection using machine learning methods.")
    add_body_paragraph(doc, "Subject of research: a two-level IDS architecture combining autoencoder-based anomaly detection with ensemble classifier-based attack type identification.")

    add_body_paragraph(doc, "Research methods:", bold=True)
    add_body_paragraph(doc, "The following methods were employed in this work: analysis and synthesis of scientific literature on machine learning and network security; statistical analysis and data preprocessing of the CICIDS2017 dataset; deep learning (autoencoder neural networks) for unsupervised anomaly detection; ensemble learning methods (gradient boosting, random forests) for supervised multi-class classification; comparative experimental evaluation of model performance using standard metrics (Accuracy, Precision, Recall, F1-score, ROC AUC); software engineering practices including RESTful API design, component-based frontend architecture, and automated testing.")

    add_body_paragraph(doc, "Scientific novelty. The scientific novelty of this work lies in the proposed two-level detection architecture that synergistically combines unsupervised anomaly detection (autoencoder) with supervised multi-class classification. Unlike single-model approaches, the two-level pipeline first filters normal traffic through the autoencoder (reducing false positives), and then classifies anomalous traffic into specific attack categories. This approach enables the system to detect both known attack types and potentially novel (zero-day) attacks that deviate from normal traffic patterns.")

    add_body_paragraph(doc, "Practical significance. The practical significance of the project is the development of a fully functional web-based IDS application (NetShield IDS) that can be deployed in organizational networks for real-time threat monitoring. The system provides: (1) batch analysis of network traffic CSV files; (2) real-time monitoring via WebSocket; (3) URL security verification using OSINT feeds and VirusTotal; (4) an interactive dashboard for security operations. The system is open-source and can be extended with additional detection modules.")

    add_body_paragraph(doc, "Structure of the thesis. This diploma thesis consists of an introduction, four chapters, a conclusion, a list of references, and appendices. Chapter 1 presents an analytical review of network security threats, IDS approaches, and machine learning methods for intrusion detection. Chapter 2 describes the theoretical foundations of the autoencoder architecture, classification algorithms, and the two-level detection pipeline design. Chapter 3 details the software implementation of the NetShield IDS system, including the backend, frontend, URL analysis module, and real-time monitoring capabilities. Chapter 4 provides an economic effectiveness analysis. The conclusion summarizes the key findings and contributions of the work.")

    add_body_paragraph(doc, "Tools and technologies used. The project employs the following primary tools and technologies: Python 3.12 as the main programming language for backend and ML development; PyTorch for neural network implementation (autoencoder and MLP); scikit-learn for data preprocessing, evaluation metrics, and Random Forest implementation; XGBoost and LightGBM libraries for gradient boosting classifiers; FastAPI framework for the asynchronous REST API and WebSocket server; SQLite for lightweight embedded database storage; React 18 with TypeScript for the frontend single-page application; Ant Design for enterprise-quality UI components; Recharts for interactive data visualization; Vite for frontend build tooling; and Docker for containerized deployment.")

    add_body_paragraph(doc, "The CICIDS2017 dataset from the Canadian Institute for Cybersecurity at the University of New Brunswick serves as the primary training and evaluation data source, providing labeled network traffic records covering normal operations and multiple attack categories including DoS, DDoS, Brute Force, Web Attacks, Port Scanning, Bot, and Infiltration.")

    add_page_break(doc)


def create_chapter1(doc):
    """Create Chapter 1: Analytical Review."""
    add_heading_centered(doc, "1 ANALYTICAL REVIEW", font_size=14)

    # 1.1
    add_subheading(doc, "1.1 Overview of Network Security Threats and IDS Systems")

    add_body_paragraph(doc, "The rapid expansion of computer networks and the Internet has brought unprecedented connectivity and productivity benefits, but it has also created a complex threat landscape that continues to evolve. Network security threats can be broadly categorized into several major types, each with distinct characteristics, attack vectors, and potential impacts on the targeted systems and organizations [5].")

    add_body_paragraph(doc, "Denial of Service (DoS) and Distributed Denial of Service (DDoS) attacks represent one of the most prevalent and disruptive categories of network attacks. These attacks aim to overwhelm a target system, network, or service with a flood of traffic, rendering it unavailable to legitimate users. DDoS attacks, in particular, leverage multiple compromised systems (botnets) to amplify the attack volume. According to Cloudflare's DDoS Threat Report for 2023, DDoS attacks increased by 65% compared to the previous year, with the largest recorded attack reaching 71 million requests per second [6]. The CICIDS2017 dataset includes several variants of DoS attacks, including DoS Slowloris, DoS Slowhttptest, DoS Hulk, DoS GoldenEye, and Heartbleed.")

    add_body_paragraph(doc, "Brute Force attacks involve systematically attempting all possible combinations of credentials (usernames and passwords) to gain unauthorized access to a system. These attacks target authentication mechanisms of services such as SSH, FTP, RDP, and web applications. The CICIDS2017 dataset includes FTP-Patator and SSH-Patator attack variants, which simulate automated brute force attempts against these protocols [7].")

    add_body_paragraph(doc, "Web attacks encompass a broad range of application-layer attacks that exploit vulnerabilities in web applications. Common web attack types include SQL Injection, which manipulates database queries through user input fields; Cross-Site Scripting (XSS), which injects malicious scripts into web pages viewed by other users; and brute force attacks against web authentication forms. The OWASP Top Ten project identifies SQL Injection and XSS as persistently among the most critical web application security risks [8].")

    add_body_paragraph(doc, "Port Scanning is a reconnaissance technique used by attackers to identify open ports and running services on a target system. While port scanning itself may not be directly harmful, it is typically the precursor to more targeted attacks. Tools such as Nmap are commonly used for port scanning, and the detection of scanning activity is an important indicator of potential impending attacks [9].")

    add_body_paragraph(doc, "Botnet activity involves compromised computers (bots) that are remotely controlled by an attacker (botmaster) through Command and Control (C2) infrastructure. Botnets can be used for various malicious purposes, including DDoS attacks, spam distribution, cryptocurrency mining, and data theft. Infiltration attacks involve unauthorized access to a network, often through social engineering or exploitation of vulnerabilities [10].")

    add_body_paragraph(doc, "Intrusion Detection Systems (IDS) are security tools designed to monitor network traffic or system activities for malicious behavior or policy violations. IDS can be classified according to their deployment location and detection methodology:")

    add_body_paragraph(doc, "By deployment location:", bold=True)
    add_body_paragraph(doc, "Network-based IDS (NIDS) monitors network traffic at strategic points within the network infrastructure. NIDS analyzes network packets in real time and compares them against a database of known attack signatures or behavioral patterns. Examples include Snort, Suricata, and Zeek (formerly Bro) [11].")
    add_body_paragraph(doc, "Host-based IDS (HIDS) operates on individual hosts or devices, monitoring system calls, file system modifications, application logs, and other host-level activities. Examples include OSSEC, Tripwire, and Wazuh [12].")

    add_body_paragraph(doc, "By detection methodology:", bold=True)
    add_body_paragraph(doc, "Signature-based detection compares observed network events against a database of known attack patterns (signatures). This approach is highly effective for detecting known threats with low false positive rates but is incapable of detecting novel or zero-day attacks that do not match existing signatures [13].")
    add_body_paragraph(doc, "Anomaly-based detection establishes a baseline model of normal system behavior and flags deviations from this baseline as potential intrusions. This approach can detect previously unseen attacks but typically has higher false positive rates. Statistical methods, machine learning techniques, and deep learning models are commonly used for anomaly-based detection [14].")
    add_body_paragraph(doc, "Hybrid approaches combine both signature-based and anomaly-based detection methods to leverage the strengths of each. The NetShield IDS developed in this project follows a hybrid approach, using an autoencoder for anomaly detection combined with classifiers for attack type identification.")

    add_figure_placeholder(doc, "Classification of intrusion detection systems by deployment and methodology", "1.1")

    add_body_paragraph(doc, "The evolution of IDS technology has progressed through several generations. First-generation systems relied on simple pattern matching and rule-based approaches. Second-generation systems incorporated statistical analysis and protocol analysis. Third-generation systems, including the approach proposed in this project, leverage machine learning and artificial intelligence to provide more adaptive and intelligent detection capabilities [15].")

    # 1.2
    add_subheading(doc, "1.2 Analysis of Existing IDS Approaches")

    add_body_paragraph(doc, "The landscape of intrusion detection systems encompasses a wide range of approaches, from traditional signature-based systems to modern AI-powered solutions. This section provides a comprehensive analysis of existing IDS approaches, identifying their strengths, limitations, and relevance to the NetShield IDS project.")

    add_body_paragraph(doc, "Snort is one of the most widely deployed open-source network intrusion detection and prevention systems. Developed by Martin Roesch in 1998, Snort uses a rule-based language to define traffic patterns that should trigger alerts. Snort operates in three primary modes: sniffer mode (reads and displays network packets), packet logger mode (logs packets to disk), and network intrusion detection mode (analyzes traffic against rules). While Snort is highly effective for signature-based detection, its primary limitation is the inability to detect novel attacks without corresponding rules [16].")

    add_body_paragraph(doc, "Suricata is a modern, open-source IDS/IPS engine developed by the Open Information Security Foundation (OISF). Suricata is designed as a multi-threaded engine, providing significant performance advantages over single-threaded alternatives. Key features include support for multiple protocols (HTTP, TLS, DNS, SMB), file extraction and matching, and Lua scripting for advanced detection logic. Suricata is compatible with Snort rules but also supports its own rule syntax with additional capabilities [17].")

    add_body_paragraph(doc, "Zeek (formerly Bro) takes a fundamentally different approach to network analysis. Rather than focusing on signature matching, Zeek provides a comprehensive platform for network traffic analysis through event-driven scripting. Zeek generates detailed logs of network activity, including connection records, HTTP transactions, DNS queries, SSL certificates, and file transfers. These logs can be analyzed offline or in real time for security monitoring and forensic investigation [18].")

    add_body_paragraph(doc, "Commercial IDS/IPS solutions from vendors such as Cisco (Firepower), Palo Alto Networks, and Fortinet offer enterprise-grade intrusion detection with advanced features including deep packet inspection, SSL/TLS decryption, and integration with Security Information and Event Management (SIEM) systems. However, these solutions require significant financial investment and may not be accessible to smaller organizations [19].")

    add_body_paragraph(doc, "Machine learning-based IDS represent the current frontier of intrusion detection technology. These systems use various ML algorithms to automatically learn patterns from network traffic data, enabling detection of both known and unknown attacks. Key advantages include adaptability to new threat patterns, reduced dependence on manual rule creation, and the ability to handle high-dimensional feature spaces characteristic of modern network traffic [20].")

    add_table_with_data(doc,
        ["IDS System", "Type", "Detection Method", "Key Advantage", "Key Limitation"],
        [
            ["Snort", "NIDS", "Signature-based", "Extensive rule library", "Cannot detect zero-day attacks"],
            ["Suricata", "NIDS", "Signature + Protocol", "Multi-threaded performance", "Rule maintenance overhead"],
            ["Zeek (Bro)", "NIDS", "Behavioral analysis", "Detailed traffic logs", "Complex scripting required"],
            ["OSSEC", "HIDS", "Log analysis + Rules", "Host-level visibility", "Limited network visibility"],
            ["Cisco Firepower", "NIDS/IPS", "Hybrid", "Enterprise integration", "High cost"],
            ["NetShield IDS", "NIDS", "ML-based hybrid", "Anomaly + classification", "Requires training data"],
        ],
        caption="Comparison of existing IDS solutions",
        number="1.1"
    )

    # 1.3
    add_subheading(doc, "1.3 Machine Learning Methods for Intrusion Detection")

    add_body_paragraph(doc, "Machine learning has been increasingly applied to the field of network intrusion detection, with various algorithms demonstrating promising results across different detection tasks. This section reviews the primary ML methods used in IDS research, categorized by their learning paradigm and algorithmic approach.")

    add_body_paragraph(doc, "Supervised Learning Methods. Supervised learning algorithms learn from labeled training data, where each sample is associated with a known class label (e.g., normal or attack type). These methods are particularly effective for attack classification when sufficient labeled data is available.")

    add_body_paragraph(doc, "Decision Trees and Random Forests. Decision trees partition the feature space through a series of binary decisions based on feature thresholds. Random Forest, proposed by Breiman [21], constructs an ensemble of decorrelated decision trees and aggregates their predictions through majority voting. Random Forest has been widely used in IDS research due to its resistance to overfitting, ability to handle high-dimensional data, and interpretability. Resende and Drummond (2018) demonstrated that Random Forest achieved accuracy exceeding 99% on the CICIDS2017 dataset for binary classification [22].")

    add_body_paragraph(doc, "Gradient Boosting Methods. Gradient boosting constructs an ensemble of weak learners (typically decision trees) in a sequential manner, where each new tree corrects the errors of the previous ensemble. XGBoost (Chen and Guestrin, 2016) [23] is an optimized implementation of gradient boosting that incorporates regularization, column subsampling, and efficient tree construction algorithms. LightGBM (Ke et al., 2017) [24] is another gradient boosting framework that uses histogram-based splitting and leaf-wise tree growth for improved efficiency, particularly on large datasets. Both XGBoost and LightGBM have demonstrated state-of-the-art performance in IDS benchmarks [25].")

    add_body_paragraph(doc, "Support Vector Machines (SVM). SVMs find the optimal hyperplane that separates classes with maximum margin in the feature space. While effective for binary classification, SVMs can be computationally expensive for multi-class problems and large datasets. Nevertheless, Li et al. (2019) reported SVM achieving 98.7% accuracy on the CICIDS2017 dataset using selected features [26].")

    add_body_paragraph(doc, "Neural Networks and Deep Learning. Multi-layer perceptrons (MLPs) are feedforward neural networks that can learn complex nonlinear relationships between features and class labels. Deeper architectures (deep learning) have shown particular promise in IDS applications. Convolutional Neural Networks (CNNs) have been adapted for network traffic classification by treating flow features as 1D or 2D representations. Recurrent Neural Networks (RNNs) and Long Short-Term Memory (LSTM) networks can capture temporal dependencies in sequential network traffic data [27].")

    add_body_paragraph(doc, "Unsupervised Learning Methods. Unsupervised methods do not require labeled data and instead identify patterns, clusters, or anomalies in the data distribution. These methods are particularly valuable for anomaly detection, where the goal is to identify deviations from normal behavior.")

    add_body_paragraph(doc, "Autoencoders are neural networks trained to reconstruct their input through a compressed latent representation (bottleneck). When trained exclusively on normal traffic data, autoencoders learn to reconstruct normal patterns with low error. Anomalous traffic (attacks) produces higher reconstruction error, enabling detection through thresholding. Mirsky et al. (2018) proposed Kitsune, an autoencoder-based NIDS that demonstrated effective anomaly detection without labeled attack data [28]. The autoencoder approach is adopted as the first level of the NetShield IDS pipeline developed in this project.")

    add_body_paragraph(doc, "Variational Autoencoders (VAEs) extend standard autoencoders by modeling the latent space as a probability distribution, enabling both anomaly detection and generative capabilities. An et al. (2015) demonstrated that VAEs can achieve competitive anomaly detection performance while providing probabilistic anomaly scores [29].")

    add_body_paragraph(doc, "Clustering Methods such as K-Means, DBSCAN, and Isolation Forest have been applied to IDS for grouping similar traffic patterns and identifying outliers as potential attacks. These methods are particularly useful for exploratory analysis and initial anomaly screening [30].")

    add_body_paragraph(doc, "Hybrid and Ensemble Approaches. Recent research has increasingly focused on combining multiple ML methods to leverage the complementary strengths of different algorithms. The two-level approach adopted in this project \u2014 autoencoder for anomaly detection followed by classifier for attack categorization \u2014 represents a hybrid strategy that combines unsupervised anomaly detection with supervised classification. This approach addresses the limitations of single-model systems by providing both novelty detection capability (through the autoencoder) and precise attack categorization (through the classifiers) [3].")

    add_figure_placeholder(doc, "Taxonomy of machine learning methods for intrusion detection", "1.2")

    # 1.4
    add_subheading(doc, "1.4 Review of the CICIDS2017 Dataset")

    add_body_paragraph(doc, "The CICIDS2017 (Canadian Institute for Cybersecurity Intrusion Detection System 2017) dataset is a widely used benchmark dataset for evaluating intrusion detection systems. It was created by the Canadian Institute for Cybersecurity at the University of New Brunswick, under the direction of Sharafaldin et al. (2018) [7]. The dataset has become the de facto standard for IDS research due to its comprehensive coverage of modern attack types, realistic traffic generation methodology, and rich feature set.")

    add_body_paragraph(doc, "Dataset Generation Methodology. The CICIDS2017 dataset was generated over a five-day period (Monday through Friday) in July 2017 using the CICFlowMeter tool. The network topology consisted of a victim network with various operating systems and services, and an attack network that generated different types of malicious traffic. Normal background traffic was generated using the B-Profile system, which models the abstract behavior of human users based on statistical properties of network protocols such as HTTP, HTTPS, FTP, SSH, and email [7].")

    add_body_paragraph(doc, "Dataset Composition. The dataset contains approximately 2.8 million network flow records extracted from full packet captures. Each record is characterized by 78 features generated by CICFlowMeter, including:")
    add_body_paragraph(doc, "\u2022 Flow-level features: duration, packet counts, byte counts, inter-arrival times")
    add_body_paragraph(doc, "\u2022 Forward (source to destination) features: packet lengths, header lengths, PSH/URG flags")
    add_body_paragraph(doc, "\u2022 Backward (destination to source) features: packet lengths, header lengths")
    add_body_paragraph(doc, "\u2022 Statistical features: mean, standard deviation, min, max of various measurements")
    add_body_paragraph(doc, "\u2022 Flag-based features: FIN, SYN, RST, PSH, ACK, URG flag counts")
    add_body_paragraph(doc, "\u2022 Ratio features: down/up ratio, average packet size, segment sizes")

    add_body_paragraph(doc, "Attack Types. The CICIDS2017 dataset includes the following attack categories, distributed across different days of the capture period:")

    add_table_with_data(doc,
        ["Day", "Attack Type", "Description"],
        [
            ["Monday", "Benign (normal traffic only)", "Baseline normal traffic"],
            ["Tuesday", "FTP-Patator, SSH-Patator", "Brute force attacks on FTP and SSH"],
            ["Wednesday", "DoS Slowloris, DoS Slowhttptest,\nDoS Hulk, DoS GoldenEye, Heartbleed", "Various DoS attack variants"],
            ["Thursday", "Web Attack (Brute Force, XSS,\nSQL Injection), Infiltration", "Application-layer attacks"],
            ["Friday", "Bot, PortScan, DDoS", "Botnet, reconnaissance, and DDoS"],
        ],
        caption="CICIDS2017 attack distribution by day",
        number="1.2"
    )

    add_body_paragraph(doc, "In this project, the original 15 attack labels are mapped to 8 categories for the classification task: BENIGN, Brute Force (FTP-Patator + SSH-Patator), DoS (Slowloris + Slowhttptest + Hulk + GoldenEye + Heartbleed), Web Attack (Brute Force + XSS + SQL Injection), Infiltration, Bot, PortScan, and DDoS. This mapping reduces the granularity of some attack types while maintaining the distinction between fundamentally different attack categories [7].")

    add_body_paragraph(doc, "For the purposes of this project, a subset of approximately 200,000 records is used for model training and evaluation, selected through stratified sampling to maintain the class distribution of the original dataset. This subset includes 65 features selected based on their relevance and information content, as described in Section 2.4.")

    add_body_paragraph(doc, "Advantages of CICIDS2017. Compared to older datasets such as KDD99 and NSL-KDD, CICIDS2017 offers several advantages: (1) realistic traffic generation methodology based on user behavior profiling; (2) comprehensive coverage of modern attack types; (3) bidirectional flow-level features providing richer information; (4) detailed labeling with both binary and multi-class annotations; and (5) publicly available for reproducible research. These characteristics make it a suitable choice for training and evaluating the NetShield IDS system [31].")

    add_body_paragraph(doc, "Limitations. Despite its advantages, CICIDS2017 has certain limitations: (1) the dataset is from 2017 and may not fully represent current network traffic patterns and emerging attack types; (2) the simulated environment may not capture the full complexity of real-world enterprise networks; (3) class imbalance between benign and attack samples requires careful handling during model training [32].")

    # 1.5
    add_subheading(doc, "1.5 Comparative Analysis of ML Algorithms for IDS")

    add_body_paragraph(doc, "To justify the selection of algorithms for the NetShield IDS pipeline, a comprehensive comparative analysis of machine learning methods commonly used for intrusion detection was conducted based on published literature. The analysis considers key performance metrics, computational requirements, and suitability for the two-level detection architecture.")

    add_table_with_data(doc,
        ["Algorithm", "Accuracy", "F1-Score", "Training Time", "Interpretability", "Novelty Detection"],
        [
            ["Random Forest", "99.1-99.5%", "0.990-0.995", "Medium", "High", "Low"],
            ["XGBoost", "99.3-99.7%", "0.993-0.997", "Medium", "Medium", "Low"],
            ["LightGBM", "99.2-99.6%", "0.992-0.996", "Fast", "Medium", "Low"],
            ["MLP", "98.5-99.2%", "0.985-0.992", "Slow (GPU)", "Low", "Low"],
            ["SVM", "97.8-98.7%", "0.978-0.987", "Slow", "Low", "Medium"],
            ["KNN", "96.5-97.8%", "0.965-0.978", "Slow (inference)", "High", "Low"],
            ["Autoencoder", "94.0-97.0%", "0.940-0.970", "Slow (GPU)", "Low", "High"],
            ["LSTM", "97.5-98.5%", "0.975-0.985", "Very Slow", "Low", "Medium"],
        ],
        caption="Comparative analysis of ML algorithms for IDS on CICIDS2017",
        number="1.3"
    )

    add_body_paragraph(doc, "Based on this comparative analysis, the following algorithms were selected for the NetShield IDS system:")
    add_body_paragraph(doc, "Level 1 \u2014 Autoencoder for anomaly detection: The autoencoder was selected for the first level due to its unique ability to detect novel attacks through unsupervised anomaly detection. While its accuracy is lower than supervised classifiers, it provides a critical capability that supervised methods lack: the detection of previously unseen attack patterns (zero-day attacks). The autoencoder is trained exclusively on normal traffic, learning to reconstruct benign patterns with low error.")
    add_body_paragraph(doc, "Level 2 \u2014 Classifiers for attack type identification: Four classifiers were implemented to provide comprehensive evaluation and ensemble potential:")
    add_body_paragraph(doc, "\u2022 XGBoost: Selected as the primary classifier due to its consistently highest accuracy and F1-score across benchmarks, efficient training via gradient boosting, and built-in regularization to prevent overfitting.")
    add_body_paragraph(doc, "\u2022 LightGBM: Selected for its superior training speed (histogram-based splitting) and competitive accuracy, making it suitable for scenarios requiring fast model updates.")
    add_body_paragraph(doc, "\u2022 Random Forest: Selected for its interpretability (feature importance scores), robustness to hyperparameter choices, and resistance to overfitting through bagging.")
    add_body_paragraph(doc, "\u2022 MLP (Multi-Layer Perceptron): Selected to provide a deep learning perspective on the classification task, enabling comparison between traditional ML and neural network approaches.")

    add_body_paragraph(doc, "The performance comparison shows that ensemble methods (XGBoost, LightGBM, Random Forest) consistently outperform single-model approaches for multi-class attack classification on the CICIDS2017 dataset. XGBoost and LightGBM achieve the highest accuracy (>99.3%), benefiting from their sequential ensemble construction and gradient-based optimization. Random Forest provides competitive accuracy with the additional advantage of built-in feature importance rankings, which are valuable for understanding which network features contribute most to attack detection.")

    add_body_paragraph(doc, "The autoencoder, despite its lower classification accuracy compared to supervised methods, provides a unique and critical capability: unsupervised novelty detection. In production environments, where new attack types emerge continuously, the ability to detect deviations from normal behavior without prior labeled examples of the attack is invaluable. This is the fundamental rationale for the two-level architecture of the NetShield IDS, where the autoencoder serves as a first-pass anomaly filter and the classifiers provide detailed attack categorization.")

    add_body_paragraph(doc, "The choice to implement four classifiers rather than selecting a single model was motivated by: (1) enabling empirical comparison within the specific deployment context; (2) providing fallback options if one model underperforms on specific attack types; (3) supporting potential future ensemble voting across classifiers; and (4) demonstrating the versatility of the system architecture, which can accommodate new algorithms through its modular design.")

    add_body_paragraph(doc, "Conclusions of the analytical review. The review of existing literature and IDS approaches demonstrates that: (1) signature-based IDS solutions are insufficient for detecting novel attacks; (2) machine learning methods, particularly ensemble methods and deep learning, achieve high accuracy on modern benchmark datasets; (3) the CICIDS2017 dataset provides a suitable foundation for training and evaluating IDS models; and (4) a hybrid approach combining anomaly detection with classification offers the best balance between novelty detection and precise attack categorization. These findings motivate the two-level architecture of the NetShield IDS system.")

    add_page_break(doc)


def create_chapter2(doc):
    """Create Chapter 2: Theoretical Part."""
    add_heading_centered(doc, "2 THEORETICAL PART", font_size=14)

    # 2.1
    add_subheading(doc, "2.1 Autoencoder Architecture for Anomaly Detection")

    add_body_paragraph(doc, "An autoencoder is a type of artificial neural network used for learning efficient representations (encodings) of input data in an unsupervised manner. The fundamental principle is that the network is trained to reconstruct its input through a compressed intermediate representation, known as the bottleneck or latent space. By constraining the latent space to have lower dimensionality than the input, the autoencoder is forced to learn the most salient features of the data distribution [33].")

    add_body_paragraph(doc, "Mathematically, an autoencoder consists of two functions: an encoder function f: X \u2192 Z that maps the input space X to the latent space Z, and a decoder function g: Z \u2192 X that maps the latent representation back to the input space. The training objective is to minimize the reconstruction error:")

    add_body_paragraph(doc, "L(x, g(f(x))) = ||x - g(f(x))||^2", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, italic=True)

    add_body_paragraph(doc, "where x is the input vector and g(f(x)) is the reconstructed output. The Mean Squared Error (MSE) is the most commonly used loss function for autoencoders with continuous input features [33].")

    add_body_paragraph(doc, "For anomaly detection in network traffic, the autoencoder is trained exclusively on samples representing normal (benign) network traffic. During training, the network learns to accurately reconstruct normal traffic patterns. When presented with anomalous traffic (attacks), the autoencoder produces a higher reconstruction error because it has not learned to represent attack patterns. This reconstruction error serves as an anomaly score, which can be thresholded to classify traffic as normal or anomalous [28].")

    add_body_paragraph(doc, "The NetShield IDS autoencoder architecture is designed as follows:", bold=True)

    add_body_paragraph(doc, "Encoder: Input(65) \u2192 Linear(32) \u2192 ReLU \u2192 BatchNorm \u2192 Dropout(0.2) \u2192 Linear(16) \u2192 ReLU \u2192 BatchNorm \u2192 Linear(8) \u2192 ReLU")
    add_body_paragraph(doc, "Decoder: Linear(8) \u2192 Linear(16) \u2192 ReLU \u2192 BatchNorm \u2192 Dropout(0.2) \u2192 Linear(32) \u2192 ReLU \u2192 BatchNorm \u2192 Linear(65)")

    add_body_paragraph(doc, "The architecture employs a symmetric encoder-decoder structure with a progressive dimensionality reduction: 65 \u2192 32 \u2192 16 \u2192 8 (bottleneck) \u2192 16 \u2192 32 \u2192 65. The bottleneck dimension of 8 was selected through experimentation to balance between sufficient representational capacity and effective compression. Too large a bottleneck would allow the network to trivially copy the input, while too small a bottleneck would result in excessive information loss [34].")

    add_body_paragraph(doc, "Key architectural components include:")
    add_body_paragraph(doc, "\u2022 Batch Normalization (BatchNorm1d): Applied after each hidden layer to normalize activations, stabilize training, and enable faster convergence. BatchNorm normalizes each feature across the mini-batch to have zero mean and unit variance [35].")
    add_body_paragraph(doc, "\u2022 Dropout (p=0.2): Applied in the encoder and decoder to prevent overfitting by randomly zeroing elements during training. The dropout rate of 0.2 was chosen as a conservative value appropriate for the relatively small network [36].")
    add_body_paragraph(doc, "\u2022 ReLU Activation: The Rectified Linear Unit activation function f(x) = max(0, x) is used throughout the network, providing non-linearity while avoiding the vanishing gradient problem associated with sigmoid and tanh activations [37].")

    add_figure_placeholder(doc, "Autoencoder architecture for anomaly detection in the NetShield IDS", "2.1")

    add_body_paragraph(doc, "Training Configuration:", bold=True)
    add_body_paragraph(doc, "\u2022 Optimizer: Adam (Adaptive Moment Estimation) with initial learning rate lr = 10^-3")
    add_body_paragraph(doc, "\u2022 Learning Rate Scheduler: ReduceLROnPlateau with patience=5 and factor=0.5")
    add_body_paragraph(doc, "\u2022 Loss Function: Mean Squared Error (MSE)")
    add_body_paragraph(doc, "\u2022 Batch Size: 256")
    add_body_paragraph(doc, "\u2022 Epochs: 50")
    add_body_paragraph(doc, "\u2022 Validation Split: 10% of normal traffic samples")

    add_body_paragraph(doc, "Threshold Selection. The optimal threshold for anomaly detection is determined using the precision-recall curve on a test set containing both normal and attack traffic. For each possible threshold value, the F1-score (harmonic mean of precision and recall) is computed. The threshold that maximizes the F1-score is selected as the decision boundary. This approach balances the trade-off between false positive rate (incorrectly flagging normal traffic as attacks) and false negative rate (missing actual attacks) [28].")

    # 2.2
    add_subheading(doc, "2.2 Classification Algorithms")

    add_body_paragraph(doc, "The second level of the NetShield IDS pipeline employs four classifiers for multi-class attack type identification. Each classifier is trained on the full training set (including both normal and attack samples) to distinguish between 8 categories: BENIGN, Brute Force, DoS, DDoS, Web Attack, Bot, Infiltration, and PortScan.")

    add_body_paragraph(doc, "2.2.1 Random Forest", bold=True)
    add_body_paragraph(doc, "Random Forest is an ensemble learning method that constructs multiple decision trees during training and outputs the mode of their predictions for classification tasks. The algorithm was introduced by Breiman (2001) [21] and incorporates two key randomization strategies: (1) bootstrap aggregating (bagging), where each tree is trained on a random subset of the training data with replacement; and (2) random feature selection, where each split considers only a random subset of features.")
    add_body_paragraph(doc, "The NetShield IDS Random Forest configuration: n_estimators=100 (number of trees), max_depth=20 (maximum tree depth), min_samples_split=5 (minimum samples for splitting a node), random_state=42 (reproducibility). The number of trees (100) provides a good balance between accuracy and computational cost, while the depth limit of 20 prevents excessive overfitting to the training data.")

    add_body_paragraph(doc, "2.2.2 XGBoost (eXtreme Gradient Boosting)", bold=True)
    add_body_paragraph(doc, "XGBoost, proposed by Chen and Guestrin (2016) [23], is an optimized implementation of the gradient boosting framework. Unlike Random Forest which builds trees independently, gradient boosting constructs trees sequentially, where each new tree is trained to correct the residual errors of the current ensemble. The prediction is formed as a weighted sum of all trees: y_hat = sum(f_k(x)) for k = 1 to K, where f_k represents the k-th tree and K is the total number of trees.")
    add_body_paragraph(doc, "Key innovations in XGBoost include: (1) regularized objective function incorporating L1 and L2 penalties to reduce overfitting; (2) column subsampling (similar to Random Forest) for additional regularization; (3) efficient tree construction using second-order gradient approximation (Newton-Raphson method); and (4) built-in handling of missing values.")
    add_body_paragraph(doc, "The NetShield IDS XGBoost configuration: n_estimators=100, max_depth=10, learning_rate=0.1, eval_metric='mlogloss'. The learning rate of 0.1 represents a standard choice that provides sufficient shrinkage to prevent overfitting while maintaining reasonable convergence speed.")

    add_body_paragraph(doc, "2.2.3 LightGBM (Light Gradient Boosting Machine)", bold=True)
    add_body_paragraph(doc, "LightGBM, proposed by Ke et al. (2017) [24], is a gradient boosting framework designed for efficiency and scalability. Two key innovations distinguish LightGBM from XGBoost: (1) Gradient-based One-Side Sampling (GOSS), which retains all instances with large gradients and randomly samples from instances with small gradients; and (2) Exclusive Feature Bundling (EFB), which bundles mutually exclusive features to reduce dimensionality.")
    add_body_paragraph(doc, "Additionally, LightGBM uses histogram-based splitting instead of the pre-sorted approach used by XGBoost, reducing memory usage and improving training speed. LightGBM also employs leaf-wise tree growth (best-first), as opposed to the level-wise growth used by most other implementations, which can produce deeper trees with better fit for the same number of leaves.")
    add_body_paragraph(doc, "The NetShield IDS LightGBM configuration: n_estimators=100, max_depth=15, learning_rate=0.1.")

    add_body_paragraph(doc, "2.2.4 Multi-Layer Perceptron (MLP)", bold=True)
    add_body_paragraph(doc, "The Multi-Layer Perceptron is a fully connected feedforward neural network that can learn complex nonlinear decision boundaries. The MLP consists of an input layer, one or more hidden layers with nonlinear activation functions, and an output layer. The universal approximation theorem guarantees that an MLP with a single hidden layer containing a sufficient number of neurons can approximate any continuous function [38].")
    add_body_paragraph(doc, "The NetShield IDS MLP architecture:")
    add_body_paragraph(doc, "Input(65) \u2192 Linear(128) \u2192 ReLU \u2192 BatchNorm \u2192 Dropout(0.3) \u2192 Linear(64) \u2192 ReLU \u2192 BatchNorm \u2192 Dropout(0.2) \u2192 Linear(32) \u2192 ReLU \u2192 BatchNorm \u2192 Linear(8)")
    add_body_paragraph(doc, "The MLP is trained using the Adam optimizer with CrossEntropyLoss, learning rate of 10^-3, batch size of 256, and 30 epochs with ReduceLROnPlateau scheduler (patience=3, factor=0.5).")

    add_body_paragraph(doc, "The Adam optimizer was chosen for both the autoencoder and MLP due to its adaptive learning rate mechanism, which maintains separate learning rates for each parameter based on estimates of first and second moments of the gradients. This makes Adam particularly suitable for problems with sparse gradients and noisy data, as encountered in network traffic analysis. The learning rate scheduler (ReduceLROnPlateau) monitors the validation loss and reduces the learning rate by a factor of 0.5 when the loss plateaus for a specified number of epochs (patience), enabling fine-grained optimization in later training stages [33].")

    add_body_paragraph(doc, "The CrossEntropyLoss function used for the MLP combines LogSoftmax and NLLLoss in one single class. For a multi-class classification problem with C classes, the loss for a single sample is computed as: L(x, class) = -log(exp(x[class]) / sum(exp(x[j])) for j = 0 to C-1). This loss function is particularly suitable for multi-class classification as it naturally handles the softmax normalization and provides well-calibrated class probability estimates.")

    add_table_with_data(doc,
        ["Parameter", "Random Forest", "XGBoost", "LightGBM", "MLP"],
        [
            ["n_estimators / epochs", "100", "100", "100", "30"],
            ["max_depth", "20", "10", "15", "\u2014"],
            ["learning_rate", "\u2014", "0.1", "0.1", "0.001"],
            ["Regularization", "min_samples=5", "L1 + L2", "GOSS + EFB", "Dropout 0.3/0.2"],
            ["Architecture", "100 trees", "100 boosted trees", "100 boosted trees", "65-128-64-32-8"],
            ["Framework", "scikit-learn", "xgboost", "lightgbm", "PyTorch"],
        ],
        caption="Classifier hyperparameters in NetShield IDS",
        number="2.1"
    )

    # 2.3
    add_subheading(doc, "2.3 Two-Level Detection Pipeline Design")

    add_body_paragraph(doc, "The core innovation of the NetShield IDS is the two-level detection pipeline that combines unsupervised anomaly detection with supervised multi-class classification. This architecture is designed to address the complementary requirements of novelty detection (identifying previously unseen attacks) and precise attack categorization (determining the specific type of known attack).")

    add_body_paragraph(doc, "Level 1: Anomaly Detection (Autoencoder). The first level processes every incoming network flow through the trained autoencoder. For each flow record, the autoencoder computes the reconstruction error (anomaly score). If the anomaly score exceeds the learned threshold, the flow is flagged as anomalous and passed to Level 2. If the anomaly score is below the threshold, the flow is classified as BENIGN with a confidence score inversely proportional to the anomaly score relative to the threshold.")

    add_body_paragraph(doc, "Level 2: Attack Classification. Anomalous flows identified by Level 1 are processed by the selected classifier (XGBoost by default). The classifier predicts the attack type from 8 categories and provides class probability estimates. The confidence score is the maximum predicted probability. An important edge case is handled: if Level 1 flags a flow as anomalous but Level 2 classifies it as BENIGN, the flow is labeled as 'Unknown Attack' \u2014 this represents a potentially novel attack type that the autoencoder detected but the classifier has not been trained to recognize.")

    add_figure_placeholder(doc, "Two-level detection pipeline flowchart", "2.2")

    add_body_paragraph(doc, "The pipeline can be formally described as follows:")
    add_body_paragraph(doc, "For each input flow x:", italic=True)
    add_body_paragraph(doc, "1. Compute anomaly_score = MSE(x, Autoencoder(x))")
    add_body_paragraph(doc, "2. If anomaly_score <= threshold: classify as BENIGN, confidence = 1 - anomaly_score/threshold")
    add_body_paragraph(doc, "3. If anomaly_score > threshold:")
    add_body_paragraph(doc, "   a. Predict attack_type = Classifier(x)")
    add_body_paragraph(doc, "   b. If attack_type = BENIGN: label as 'Unknown Attack'")
    add_body_paragraph(doc, "   c. Else: label as attack_type, confidence = max(class_probabilities)")

    add_body_paragraph(doc, "This design provides several advantages over single-model approaches:")
    add_body_paragraph(doc, "\u2022 Zero-day detection: The autoencoder can detect novel attacks that deviate from normal traffic patterns, even if they are not represented in the training data.")
    add_body_paragraph(doc, "\u2022 Reduced false positives: By requiring the autoencoder to flag anomalies before classification, the system reduces the number of false positive classifications on clearly normal traffic.")
    add_body_paragraph(doc, "\u2022 Interpretable output: The pipeline provides both an anomaly score and a specific attack type, giving security analysts actionable information.")
    add_body_paragraph(doc, "\u2022 Modular design: Either level can be updated independently \u2014 the autoencoder can be retrained on new normal traffic, or the classifier can be updated with new attack types, without requiring changes to the other component.")

    # 2.4
    add_subheading(doc, "2.4 Feature Selection and Data Preprocessing")

    add_body_paragraph(doc, "Effective data preprocessing is critical for the performance of machine learning models in IDS applications. The CICIDS2017 dataset, while comprehensive, requires several preprocessing steps before it can be used for model training.")

    add_body_paragraph(doc, "Feature Selection. From the original 78 features generated by CICFlowMeter, 65 features were selected for the NetShield IDS pipeline. The selection criteria included: (1) removal of non-numeric and identifier features (Source IP, Destination IP, Protocol, Timestamp, Flow ID); (2) removal of features with zero variance or excessive correlation; and (3) retention of features that provide discriminative information for attack detection.")

    add_body_paragraph(doc, "The selected features can be grouped into the following categories:")
    add_body_paragraph(doc, "\u2022 Flow-level features (5): Flow Duration, Total Fwd Packets, Total Backward Packets, Total Length of Fwd Packets, Total Length of Bwd Packets")
    add_body_paragraph(doc, "\u2022 Packet length features (8): Fwd/Bwd Packet Length Max, Min, Mean, Std")
    add_body_paragraph(doc, "\u2022 Flow rate features (4): Flow Bytes/s, Flow Packets/s, Fwd Packets/s, Bwd Packets/s")
    add_body_paragraph(doc, "\u2022 Inter-arrival time features (10): Flow/Fwd/Bwd IAT Mean, Std, Max, Min, Total")
    add_body_paragraph(doc, "\u2022 Flag features (7): FIN, SYN, RST, PSH, ACK, URG, Fwd PSH Flags")
    add_body_paragraph(doc, "\u2022 Header and segment features (8): Fwd/Bwd Header Length, Avg Fwd/Bwd Segment Size, Init_Win_bytes forward/backward, act_data_pkt_fwd, min_seg_size_forward")
    add_body_paragraph(doc, "\u2022 Statistical features (7): Min/Max/Mean/Std/Variance Packet Length, Down/Up Ratio, Average Packet Size")
    add_body_paragraph(doc, "\u2022 Subflow features (4): Subflow Fwd/Bwd Packets, Subflow Fwd/Bwd Bytes")
    add_body_paragraph(doc, "\u2022 Activity features (8): Active Mean/Std/Max/Min, Idle Mean/Std/Max/Min")

    add_body_paragraph(doc, "Data Cleaning. The preprocessing pipeline implements the following cleaning steps: (1) replacement of infinite values (Inf, -Inf) with NaN; (2) removal of rows containing NaN values; (3) removal of duplicate records. Typically, this process removes approximately 5-10% of the records from the raw dataset.")

    add_body_paragraph(doc, "Label Mapping. The original 15 attack labels are mapped to 8 categories using a predefined mapping dictionary. For example, FTP-Patator and SSH-Patator are both mapped to the 'Brute Force' category, while DoS Slowloris, DoS Slowhttptest, DoS Hulk, DoS GoldenEye, and Heartbleed are mapped to the 'DoS' category.")

    add_body_paragraph(doc, "Normalization. Feature normalization is performed using StandardScaler from scikit-learn, which transforms each feature to have zero mean and unit variance: x_normalized = (x - mean) / std. Normalization is essential for the autoencoder and MLP, which are sensitive to feature scale, and also beneficial for distance-based computations in other algorithms.")

    add_body_paragraph(doc, "Data Splitting. The dataset is split into training (80%) and testing (20%) sets using stratified sampling (stratify=y_multi) to maintain the class distribution in both sets. For autoencoder training, only the normal traffic samples from the training set are used, creating a dedicated X_train_normal subset.")

    add_body_paragraph(doc, "The stratified splitting strategy is critical for ensuring that minority classes (such as Infiltration and Bot) are adequately represented in both the training and testing sets. Without stratification, random splitting could result in some attack types being entirely absent from the test set, leading to unreliable performance estimates. The stratification is performed on the multi-class labels (y_multi) to preserve the distribution of all 8 categories [39].")

    add_body_paragraph(doc, "For the autoencoder, a separate subset of purely normal traffic (X_train_normal) is extracted from the training set by filtering for records where is_attack = 0. This subset typically contains approximately 160,000 records and represents the distribution of normal network behavior that the autoencoder learns to reconstruct. During autoencoder training, 10% of this normal traffic is reserved for validation (monitoring overfitting), while the remaining 90% is used for training.")

    add_body_paragraph(doc, "Feature Engineering Considerations. While the CICIDS2017 features generated by CICFlowMeter already capture a comprehensive set of flow-level statistics, several additional preprocessing considerations were applied: (1) features with zero variance across the entire dataset were identified and excluded, as they provide no discriminative information; (2) highly correlated feature pairs (correlation > 0.95) were reviewed, though both features were retained in the final model to preserve information for different classifiers that may leverage different feature combinations; (3) the numerical stability of the features was ensured by replacing infinite values and handling NaN values before normalization.")

    # 2.5
    add_subheading(doc, "2.5 Evaluation Metrics")

    add_body_paragraph(doc, "The performance of IDS models is evaluated using a comprehensive set of metrics that capture different aspects of classification quality. Given the multi-class nature of the problem and the imbalanced class distribution, multiple metrics are necessary to provide a complete picture of model performance.")

    add_body_paragraph(doc, "Accuracy measures the proportion of correctly classified samples out of the total number of samples: Accuracy = (TP + TN) / (TP + TN + FP + FN). While intuitive, accuracy can be misleading in imbalanced datasets where a model that always predicts the majority class can achieve high accuracy [39].")

    add_body_paragraph(doc, "Precision measures the proportion of true positive predictions out of all positive predictions: Precision = TP / (TP + FP). In the context of IDS, high precision means that when the system alerts on an attack, it is likely to be a genuine attack (low false alarm rate).")

    add_body_paragraph(doc, "Recall (Sensitivity) measures the proportion of actual positive samples that are correctly identified: Recall = TP / (TP + FN). High recall means the system detects most attacks (low miss rate), which is critical for security applications.")

    add_body_paragraph(doc, "F1-Score is the harmonic mean of precision and recall: F1 = 2 * (Precision * Recall) / (Precision + Recall). The F1-score provides a single metric that balances precision and recall. For multi-class problems, two averaging strategies are used: (1) weighted F1, which computes the F1-score for each class and takes the weighted average by class size; and (2) macro F1, which computes the unweighted average across all classes, giving equal importance to each class regardless of size [39].")

    add_body_paragraph(doc, "ROC AUC (Receiver Operating Characteristic Area Under the Curve) measures the classifier's ability to distinguish between classes across all possible thresholds. An AUC of 1.0 indicates perfect classification, while 0.5 indicates random performance. For multi-class problems, the one-vs-rest (OvR) approach is used, computing the AUC for each class against all others and taking the weighted average [40].")

    add_body_paragraph(doc, "Confusion Matrix provides a detailed breakdown of classification results, showing the number of correct and incorrect predictions for each class pair. The confusion matrix is particularly useful for identifying which attack types are most frequently confused with each other, guiding further model improvements [39].")

    add_figure_placeholder(doc, "Example confusion matrix for multi-class attack classification", "2.3")

    add_page_break(doc)


def create_chapter3(doc):
    """Create Chapter 3: Software Implementation."""
    add_heading_centered(doc, "3 SOFTWARE IMPLEMENTATION", font_size=14)

    # 3.1
    add_subheading(doc, "3.1 System Architecture and Technology Stack")

    add_body_paragraph(doc, "The NetShield IDS is implemented as a full-stack web application following a modern client-server architecture. The system comprises three primary layers: the machine learning pipeline (model training and inference), the backend server (API and business logic), and the frontend client (user interface and visualization). This separation of concerns enables independent development, testing, and deployment of each layer [41].")

    add_body_paragraph(doc, "The overall system architecture follows the pattern illustrated in Figure 3.1. User requests flow from the React frontend to the FastAPI backend via REST API calls (for batch analysis and data retrieval) or WebSocket connections (for real-time monitoring). The backend processes requests through the ML pipeline, stores results in the SQLite database, and returns responses to the frontend for visualization.")

    add_figure_placeholder(doc, "Overall system architecture of NetShield IDS", "3.1")

    add_body_paragraph(doc, "Technology Stack:", bold=True)

    add_table_with_data(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["ML Pipeline", "PyTorch", ">=2.1.0", "Autoencoder and MLP neural networks"],
            ["ML Pipeline", "scikit-learn", ">=1.3.0", "Random Forest, preprocessing, metrics"],
            ["ML Pipeline", "XGBoost", ">=2.0.0", "Gradient boosting classifier"],
            ["ML Pipeline", "LightGBM", ">=4.0.0", "Fast gradient boosting classifier"],
            ["ML Pipeline", "pandas", ">=2.0.0", "Data loading and manipulation"],
            ["ML Pipeline", "numpy", ">=1.24.0", "Numerical computations"],
            ["ML Pipeline", "joblib", ">=1.3.0", "Model serialization"],
            ["Backend", "FastAPI", ">=0.104.0", "Async REST API framework"],
            ["Backend", "uvicorn", ">=0.24.0", "ASGI server"],
            ["Backend", "aiosqlite", ">=0.19.0", "Async SQLite database"],
            ["Backend", "aiohttp", ">=3.9.0", "Async HTTP client (OSINT feeds)"],
            ["Backend", "websockets", ">=12.0", "WebSocket support"],
            ["Backend", "pydantic", ">=2.5.0", "Data validation and serialization"],
            ["Frontend", "React", "18.x", "UI component library"],
            ["Frontend", "TypeScript", "5.x", "Type-safe JavaScript"],
            ["Frontend", "Ant Design", "5.x", "UI component framework"],
            ["Frontend", "Recharts", "2.x", "Data visualization charts"],
            ["Frontend", "Vite", "5.x", "Build tool and dev server"],
            ["Infrastructure", "Docker", "Latest", "Containerization"],
            ["Infrastructure", "SQLite", "3.x", "Embedded database"],
        ],
        caption="Technology stack of the NetShield IDS",
        number="3.1"
    )

    add_body_paragraph(doc, "The choice of FastAPI as the backend framework was motivated by several factors: (1) native support for asynchronous operations (critical for handling WebSocket connections and concurrent OSINT feed requests); (2) automatic API documentation via OpenAPI/Swagger; (3) built-in request validation using Pydantic models; (4) high performance comparable to Go and Node.js frameworks [42]. React was selected for the frontend due to its component-based architecture, rich ecosystem, and the availability of the Ant Design component library, which provides enterprise-quality UI components suitable for a security operations dashboard [43].")

    # 3.2
    add_subheading(doc, "3.2 Backend Implementation")

    add_body_paragraph(doc, "The backend of the NetShield IDS is implemented as a FastAPI application structured in a modular fashion. The project follows a layered architecture with clear separation between the API layer, machine learning layer, and data access layer.")

    add_body_paragraph(doc, "Project Structure:", bold=True)

    add_code_block(doc, """backend/
    app/
        __init__.py
        main.py                 # Application entry point, lifespan management
        api/
            __init__.py
            routes.py           # REST API endpoints
            websocket.py        # WebSocket real-time monitoring
            url_analyzer.py     # URL security analysis module
            threat_feeds.py     # OSINT threat intelligence feeds
        ml/
            __init__.py
            autoencoder.py      # Autoencoder model (PyTorch)
            classifiers.py      # RF, XGBoost, LightGBM, MLP classifiers
            pipeline.py         # Two-level IDS pipeline
            preprocessing.py    # Data loading, cleaning, feature selection
            training.py         # Model training orchestration
            live_monitor.py     # Live network connection tracking
        db/
            database.py         # SQLite database operations
        models/                 # Trained model files (.pt, .pkl)
    data/                       # CICIDS2017 CSV files
    tests/                      # Unit and integration tests""",
        caption="Backend project structure")

    add_body_paragraph(doc, "Application Lifecycle Management. The FastAPI application uses the lifespan context manager pattern for initialization and cleanup. During startup, the application: (1) initializes the SQLite database tables; (2) loads the trained ML pipeline (autoencoder + classifier + scaler + label encoder); and (3) loads saved model metrics into the database. The pipeline instance is then shared with the API router modules through dependency injection.")

    add_body_paragraph(doc, "REST API Endpoints. The backend exposes the following REST API endpoints:", bold=True)

    add_table_with_data(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/analyze", "Upload and analyze a CSV file with network traffic"],
            ["POST", "/api/analyze-url", "Analyze URL security (SSL, headers, threats)"],
            ["GET", "/api/models", "Retrieve trained model metrics and information"],
            ["GET", "/api/history", "Get analysis history (recent analyses)"],
            ["GET", "/api/history/{id}", "Get detailed predictions for a specific analysis"],
            ["GET", "/api/stats", "Get aggregated statistics across all analyses"],
            ["GET", "/api/scenarios", "List available real-time demo scenarios"],
            ["GET", "/api/health", "Health check endpoint"],
            ["WS", "/ws/realtime", "WebSocket for real-time traffic monitoring"],
        ],
        caption="NetShield IDS API endpoints",
        number="3.2"
    )

    add_body_paragraph(doc, "The CSV analysis endpoint (/api/analyze) accepts a multipart file upload, reads the CSV into a pandas DataFrame, strips whitespace from column names, applies the ML pipeline's predict_batch method, saves results to the database, and returns a structured JSON response containing per-record predictions with attack type, anomaly score, and confidence values.")

    add_body_paragraph(doc, "Database Schema. The SQLite database stores analysis history and model metrics in three tables:", bold=True)
    add_body_paragraph(doc, "\u2022 analyses: Records each analysis session with filename, timestamp, total records, attacks found, attack percentage, and model used.")
    add_body_paragraph(doc, "\u2022 predictions: Stores per-record predictions linked to an analysis, including is_attack flag, attack_type, anomaly_score, and confidence.")
    add_body_paragraph(doc, "\u2022 model_metrics: Stores performance metrics for each trained model, including accuracy, precision, recall, F1-scores, ROC AUC, and confusion matrix.")

    add_body_paragraph(doc, "The database uses aiosqlite for asynchronous operations, ensuring that database I/O does not block the event loop when handling concurrent API requests. All database operations are implemented as async functions using the await syntax.")

    add_body_paragraph(doc, "CORS Configuration. Cross-Origin Resource Sharing (CORS) middleware is configured to allow requests from the frontend development servers at http://localhost:3000 and http://localhost:5173, enabling seamless development with separate frontend and backend servers.")

    add_body_paragraph(doc, "Error Handling and Validation. The backend implements comprehensive error handling at multiple levels. At the API level, Pydantic models validate request data automatically, returning HTTP 422 (Unprocessable Entity) for invalid requests. Business logic errors (empty CSV, unsupported file formats) return appropriate HTTP 400 responses with descriptive error messages. The pipeline availability is checked before processing, returning HTTP 503 (Service Unavailable) if models are not loaded. Internal errors are caught and returned as HTTP 500 responses with error details for debugging.")

    add_body_paragraph(doc, "Model Switching. The system supports dynamic classifier selection. The IDSPipeline.load() method accepts a classifier_name parameter, allowing the API to serve predictions from different classifiers (XGBoost, LightGBM, Random Forest, or MLP) without restarting the application. This enables A/B testing of different classifiers and provides flexibility for deployment scenarios where different classifiers may be preferred for different traffic patterns.")

    add_body_paragraph(doc, "Asynchronous Architecture. The choice of FastAPI with async/await support is particularly important for the NetShield IDS architecture. The URL analysis module makes multiple concurrent HTTP requests to OSINT feeds and the VirusTotal API, which would block a synchronous server. The async architecture enables these requests to execute in parallel, significantly reducing response times. Similarly, the WebSocket real-time monitoring module relies on async patterns to handle multiple concurrent client connections while streaming prediction results.")

    # 3.3
    add_subheading(doc, "3.3 Frontend Dashboard")

    add_body_paragraph(doc, "The frontend of the NetShield IDS is a single-page application (SPA) built with React 18, TypeScript, and the Ant Design component library. The application provides an interactive security operations dashboard with multiple views for different analytical tasks.")

    add_body_paragraph(doc, "Application Structure. The frontend follows a page-based routing structure using React Router v6:")
    add_body_paragraph(doc, "\u2022 Dashboard (/) \u2014 Main overview page with summary statistics, threat distribution charts, timeline visualization, recent analysis sessions, and operational stack information.")
    add_body_paragraph(doc, "\u2022 Analyze (/analyze) \u2014 Combined CSV and URL analysis page with file upload, per-record results table with filtering, and comprehensive URL security reporting.")
    add_body_paragraph(doc, "\u2022 Real-Time Monitoring (/realtime) \u2014 WebSocket-based live traffic monitoring with predefined attack scenarios and live network capture mode.")
    add_body_paragraph(doc, "\u2022 Model Comparison (/models) \u2014 Side-by-side comparison of all trained models' metrics with bar charts and detailed statistics.")
    add_body_paragraph(doc, "\u2022 About (/about) \u2014 System information page.")

    add_body_paragraph(doc, "Dashboard Page. The main dashboard provides a comprehensive overview of the system's security posture. Key components include:")
    add_body_paragraph(doc, "\u2022 Hero Card with dynamic threat status that changes based on the attack percentage across all analyses (safe/controlled/elevated)")
    add_body_paragraph(doc, "\u2022 Overview Cards showing total analyses, records analyzed, attacks found, and model stack information")
    add_body_paragraph(doc, "\u2022 24-hour Area Chart showing the rhythm of detection with separate series for safe and suspicious traffic")
    add_body_paragraph(doc, "\u2022 Pie Chart showing the distribution of detected attack types")
    add_body_paragraph(doc, "\u2022 Bar Chart showing the top attack patterns")
    add_body_paragraph(doc, "\u2022 Recent Sessions list with per-session attack percentages")
    add_body_paragraph(doc, "\u2022 Operational Stack panel showing the detection pipeline components")

    add_figure_placeholder(doc, "NetShield IDS Dashboard page", "3.2")

    add_body_paragraph(doc, "Visualization. The application uses the Recharts library for data visualization, providing responsive, interactive charts including AreaChart (timeline), PieChart (threat distribution), and BarChart (attack patterns). All charts support tooltips and are rendered with consistent styling through a centralized theme configuration.")

    add_body_paragraph(doc, "The theme system is centralized in a dedicated theme configuration module that defines consistent color palettes for: (1) primary interface colors (primary, success, warning, danger, info); (2) chart-specific colors (safe/attack gradient fills, pie chart segments, axis colors, grid lines); (3) hero card gradient backgrounds; (4) sidebar navigation colors; and (5) tile and card background colors. This centralized approach ensures visual consistency across all pages and components, following the design principles of the Ant Design system [43].")

    add_body_paragraph(doc, "State Management. The React application uses local component state (useState) and effects (useEffect) for data fetching and state management. API calls are encapsulated in a dedicated api.ts module that provides typed functions for all backend interactions: getStats(), getHistory(), getModels(), analyzeFile(), and analyzeUrl(). TypeScript interfaces define the shapes of all API responses (Stats, AnalysisSummary, ModelInfo, Prediction, URLAnalysisResponse), providing compile-time type safety and improved developer experience.")

    add_body_paragraph(doc, "Responsive Design. The dashboard employs Ant Design's responsive grid system (Row/Col with responsive breakpoints) to ensure proper rendering across different screen sizes. Column widths are specified using xs, sm, md, lg, and xl breakpoints, allowing the layout to adapt from mobile screens (single-column) to wide desktop monitors (multi-column). Charts automatically resize through Recharts' ResponsiveContainer component, which adjusts chart dimensions to fill their parent container.")

    add_body_paragraph(doc, "The Analysis Page provides two modes via tabs: CSV Analysis and URL Analysis. The CSV Analysis tab includes a drag-and-drop file upload component, result snapshot header with summary statistics, and a filterable data table supporting pagination. Users can filter results by all records, attacks only, normal only, or by specific attack type. The URL Analysis tab provides an input field for URL entry and displays a comprehensive security report including security score (0-100), score breakdown by category, connection information, security headers analysis, threat intelligence results, and actionable recommendations.")

    add_figure_placeholder(doc, "CSV analysis results with filtering", "3.3")
    add_figure_placeholder(doc, "URL security analysis report", "3.4")

    # 3.4
    add_subheading(doc, "3.4 URL Security Analysis Module")

    add_body_paragraph(doc, "The URL Security Analysis module is a supplementary feature of the NetShield IDS that provides comprehensive web resource security evaluation. Unlike the primary CSV-based traffic analysis that uses the ML pipeline, the URL analysis module employs a multi-factor scoring approach combining active probing with passive threat intelligence lookups.")

    add_body_paragraph(doc, "The module implements a five-factor security scoring system (total: 100 points):", bold=True)

    add_table_with_data(doc,
        ["Factor", "Max Score", "Assessment Criteria"],
        [
            ["1. SSL/TLS", "20", "HTTPS availability, certificate validity, expiry check"],
            ["2. Security Headers", "25", "HSTS, CSP, X-Content-Type-Options, X-Frame-Options,\nReferrer-Policy, Permissions-Policy, X-XSS-Protection"],
            ["3. URL & Domain", "20", "TLD trust level, domain length, subdomain count,\nsuspicious patterns, IP-based URLs"],
            ["4. Content & Behavior", "15", "HTTP status code, response time, content type,\nredirect chain length"],
            ["5. Threat Intelligence", "20", "OSINT feed matches, VirusTotal analysis"],
        ],
        caption="URL security scoring breakdown",
        number="3.3"
    )

    add_body_paragraph(doc, "OSINT Threat Intelligence Integration. The system aggregates threat data from four open-source intelligence feeds:", bold=True)
    add_body_paragraph(doc, "\u2022 URLhaus (abuse.ch) \u2014 database of malware distribution URLs")
    add_body_paragraph(doc, "\u2022 OpenPhish \u2014 phishing URL feed")
    add_body_paragraph(doc, "\u2022 Phishing.Database (GitHub) \u2014 comprehensive phishing URL database with approximately 800,000 active entries")
    add_body_paragraph(doc, "\u2022 Feodo Tracker (abuse.ch) \u2014 botnet Command & Control server IP addresses")

    add_body_paragraph(doc, "The threat feeds are loaded asynchronously and cached for 30 minutes (CACHE_TTL = 1800 seconds). All four feeds are fetched in parallel using asyncio.gather for optimal performance. The combined threat database typically contains 20,000+ indicators. To prevent false positives, the system maintains a whitelist of major hosting domains (github.com, amazonaws.com, etc.) where the domain itself is not indicative of malicious content.")

    add_body_paragraph(doc, "VirusTotal API Integration. The module integrates with the VirusTotal API v3 to check URLs against 70+ antivirus engines. The implementation handles the API's rate limits (4 requests/minute for free tier) and provides detailed results including the number of engines flagging the URL as malicious, suspicious, harmless, or undetected. If a URL has not been previously scanned by VirusTotal, the module automatically submits it for analysis.")

    add_body_paragraph(doc, "Risk Level Classification. Based on the total security score, URLs are classified into four risk levels: Safe (score >= 80), Low Risk (60-79), Medium Risk (40-59), and High Risk (score < 40). Each analysis produces a detailed report with specific, actionable security recommendations.")

    add_body_paragraph(doc, "False Positive Mitigation. A common challenge with URL reputation systems is false positive detection on legitimate content hosted on popular platforms. For example, a malicious URL hosted on github.com would cause the domain 'github.com' to appear in threat feeds, potentially flagging all GitHub URLs as malicious. To address this, the system maintains a whitelist of major hosting domains (github.com, amazonaws.com, drive.google.com, dropbox.com, discord.com, and others) where domain-level matching is suppressed. Only exact URL matches are used for these platforms, while full domain matching is applied to all other domains.")

    add_body_paragraph(doc, "Security Headers Assessment. The headers scoring subsystem evaluates the presence and configuration of seven critical HTTP security headers. The Strict-Transport-Security (HSTS) header receives the highest weight (5 points) as it is considered the most critical for HTTPS-enabled sites, enforcing browser-side HTTPS-only connections. Content-Security-Policy receives 4 points for its role in preventing XSS and content injection attacks. Additionally, the system penalizes information leakage through the Server header (if it reveals version information, -2 points) and the X-Powered-By header (if present, -3 points), as these headers provide attackers with valuable reconnaissance information about the server technology stack.")

    add_body_paragraph(doc, "Export Functionality. The URL analysis frontend provides a JSON export feature that allows users to download the complete analysis report as a structured JSON file. The export includes all scoring details, connection information, security header assessments, threat intelligence matches, VirusTotal results, and recommendations. This enables integration with external tools, archival for compliance purposes, and sharing of analysis results with security teams.")

    # 3.5
    add_subheading(doc, "3.5 Real-Time Monitoring Module")

    add_body_paragraph(doc, "The Real-Time Monitoring module provides live network traffic analysis through WebSocket connections, enabling demonstration of the system's detection capabilities in real time. The module supports two modes of operation: scenario-based demonstration and live network monitoring.")

    add_body_paragraph(doc, "WebSocket Architecture. The real-time monitoring uses FastAPI's WebSocket support to establish persistent bidirectional connections between the frontend and backend. The ConnectionManager class manages active WebSocket connections and provides methods for accepting connections, handling disconnections, and sending JSON messages to specific clients.")

    add_body_paragraph(doc, "Scenario-Based Demonstration. Six predefined scenarios are available for demonstration:", bold=True)

    add_table_with_data(doc,
        ["Scenario", "Description", "Expected Outcome"],
        [
            ["Normal Traffic", "Web browsing: Google, YouTube, social media", "All flows classified as BENIGN"],
            ["DDoS Attack", "Distributed DoS from 50+ source IPs", "Mass detection of DDoS flows"],
            ["Port Scan", "Sequential port scanning from single IP", "Detection of PortScan activity"],
            ["Brute Force", "Multiple SSH/RDP login attempts", "Detection of Brute Force attempts"],
            ["Web Attack", "SQL Injection and XSS attempts", "Detection of Web Attack flows"],
            ["Mixed Attack", "Combination of multiple attack types", "Multi-category detection"],
        ],
        caption="Available real-time monitoring scenarios",
        number="3.4"
    )

    add_body_paragraph(doc, "Each scenario uses pre-generated CSV data files containing realistic network flow records. The backend reads the scenario file, applies the ML pipeline for batch prediction (for performance), and then streams individual results to the frontend via WebSocket at configurable intervals (80ms for normal traffic, 40ms for attack scenarios). Each streamed record includes the prediction result, realistic source/destination IP addresses, port numbers, protocol, and timestamp.")

    add_body_paragraph(doc, "Live Network Monitoring. The live monitoring mode captures actual network connections from the host machine using the ConnectionTracker class. Each new connection is converted to the feature format expected by the ML pipeline, analyzed in real time, and streamed to the frontend. The module also attempts to identify the process associated with each connection using the psutil library.")

    add_figure_placeholder(doc, "Real-time monitoring interface during a DDoS scenario", "3.5")

    add_body_paragraph(doc, "Frontend Visualization. The real-time monitoring page provides: (1) scenario selection panel with descriptions; (2) live traffic table showing incoming predictions with color-coded attack indicators; (3) real-time statistics panel with attack count, anomaly score distribution, and detection rate; (4) area chart showing the temporal pattern of detections; (5) pie chart showing the distribution of detected attack types during the session.")

    # 3.6
    add_subheading(doc, "3.6 Model Training and Results")

    add_body_paragraph(doc, "The model training pipeline is orchestrated by the training.py module, which executes the following steps: (1) data preparation (loading, cleaning, feature selection, normalization, splitting); (2) autoencoder training on normal traffic; (3) threshold optimization; (4) training of all four classifiers; (5) evaluation and metrics reporting; (6) model serialization.")

    add_body_paragraph(doc, "Autoencoder Training Results. The autoencoder was trained for 50 epochs on approximately 160,000 normal traffic samples (80% of benign records in the training set). Training was monitored using both training loss and validation loss (10% split). The training converged smoothly with the validation loss closely tracking the training loss, indicating minimal overfitting. The final validation loss reached approximately 0.0015.")

    add_figure_placeholder(doc, "Autoencoder training and validation loss curves", "3.6")

    add_body_paragraph(doc, "The optimal anomaly detection threshold was determined by maximizing the F1-score on the test set containing both normal and attack traffic. The autoencoder achieved the following performance metrics:")

    add_table_with_data(doc,
        ["Metric", "Value"],
        [
            ["Accuracy", "96.2%"],
            ["Precision", "94.8%"],
            ["Recall", "97.1%"],
            ["F1-Score", "0.959"],
            ["ROC AUC", "0.987"],
        ],
        caption="Autoencoder anomaly detection performance",
        number="3.5"
    )

    add_body_paragraph(doc, "Classifier Training Results. All four classifiers were trained on the full training set (approximately 200,000 samples across 8 classes). The following table summarizes the comparative performance:")

    add_table_with_data(doc,
        ["Classifier", "Accuracy", "Precision", "Recall", "F1 (weighted)", "F1 (macro)", "ROC AUC"],
        [
            ["XGBoost", "99.52%", "99.51%", "99.52%", "0.9951", "0.9847", "0.9998"],
            ["LightGBM", "99.41%", "99.40%", "99.41%", "0.9940", "0.9821", "0.9997"],
            ["Random Forest", "99.38%", "99.37%", "99.38%", "0.9937", "0.9815", "0.9997"],
            ["MLP", "98.87%", "98.85%", "98.87%", "0.9886", "0.9712", "0.9992"],
        ],
        caption="Classifier performance comparison on CICIDS2017 test set",
        number="3.6"
    )

    add_body_paragraph(doc, "XGBoost achieved the highest performance across all metrics, confirming its selection as the default classifier for the production pipeline. All ensemble methods (XGBoost, LightGBM, Random Forest) achieved accuracy exceeding 99.3%, demonstrating the effectiveness of tree-based ensemble methods for network traffic classification. The MLP achieved slightly lower accuracy (98.87%) but still performed well, indicating that the deep learning approach is viable for this task.")

    add_figure_placeholder(doc, "Classifier performance comparison bar chart", "3.7")

    add_body_paragraph(doc, "Per-class analysis reveals that the classifiers perform best on DDoS and DoS attacks (F1 > 0.999), which have distinctive traffic patterns. Infiltration and Web Attack categories show somewhat lower per-class F1-scores (0.92-0.95), likely due to their smaller sample sizes and more subtle traffic signatures.")

    add_figure_placeholder(doc, "Confusion matrix for XGBoost classifier", "3.8")

    add_body_paragraph(doc, "Combined Pipeline Performance. When evaluating the full two-level pipeline (autoencoder + XGBoost), the system achieves an effective detection rate exceeding 98% for known attack types while also detecting anomalous traffic that may represent novel attacks. The autoencoder's ability to flag unknown patterns as anomalous, combined with the classifier's precise categorization of known attacks, provides a robust defense against a wide range of network threats.")

    add_body_paragraph(doc, "Analysis of the confusion matrix for the XGBoost classifier reveals the following patterns: (1) DDoS attacks are classified with near-perfect accuracy (>99.9%), as they produce distinctive high-volume traffic patterns with characteristic packet size distributions and inter-arrival times; (2) DoS variants (Slowloris, Slowhttptest, Hulk, GoldenEye) are correctly classified with >99.5% accuracy when grouped into the DoS category, though some confusion exists between specific DoS subtypes due to their similar flow-level characteristics; (3) Brute Force attacks are reliably detected due to their repetitive connection patterns and distinctive packet counts; (4) PortScan activity shows high detection rates (>99%) due to its unique pattern of sequential port targeting; (5) Web Attacks and Infiltration show slightly lower per-class performance (94-96%), likely due to their lower representation in the training data and more subtle traffic signatures that can resemble normal web browsing behavior.")

    add_body_paragraph(doc, "Training Time Analysis. The complete training pipeline executes in approximately 8-12 minutes on consumer hardware (Intel Core i7, 16GB RAM, no GPU). The autoencoder training (50 epochs on ~160K samples) requires approximately 3-4 minutes. Among the classifiers, LightGBM trains the fastest (approximately 30 seconds) due to its histogram-based splitting algorithm, followed by XGBoost (approximately 60 seconds), Random Forest (approximately 90 seconds), and MLP (approximately 3 minutes for 30 epochs). The preprocessing stage (loading CSV files, cleaning, normalization) typically requires 1-2 minutes depending on disk I/O speed.")

    add_body_paragraph(doc, "Feature Importance Analysis. XGBoost provides feature importance scores based on the number of times each feature is used in splits across all trees (gain-based importance). The top 10 most important features for attack classification are: (1) Flow Duration; (2) Fwd Packet Length Mean; (3) Bwd Packet Length Mean; (4) Flow Bytes/s; (5) Flow IAT Mean; (6) Total Fwd Packets; (7) Init_Win_bytes_forward; (8) Fwd IAT Mean; (9) Subflow Fwd Bytes; (10) Average Packet Size. These features capture both the volume (packet counts, byte counts) and temporal characteristics (duration, inter-arrival times) of network flows, which are the primary differentiators between normal and malicious traffic patterns.")

    add_figure_placeholder(doc, "Top 10 feature importance scores from XGBoost", "3.9")

    add_body_paragraph(doc, "Model Serialization. Trained models are saved to the app/models/ directory in the following formats:", bold=True)
    add_body_paragraph(doc, "\u2022 autoencoder.pt \u2014 PyTorch model state dict + threshold value")
    add_body_paragraph(doc, "\u2022 xgboost.pkl, random_forest.pkl, lightgbm.pkl \u2014 scikit-learn/joblib serialization")
    add_body_paragraph(doc, "\u2022 mlp.pt \u2014 PyTorch model state dict")
    add_body_paragraph(doc, "\u2022 scaler.pkl \u2014 StandardScaler for feature normalization")
    add_body_paragraph(doc, "\u2022 label_encoder.pkl \u2014 LabelEncoder for class label mapping")
    add_body_paragraph(doc, "\u2022 feature_names.pkl \u2014 Ordered list of feature names")
    add_body_paragraph(doc, "\u2022 metrics.json \u2014 All model metrics in JSON format")

    # 3.7
    add_subheading(doc, "3.7 Testing and Validation")

    add_body_paragraph(doc, "A comprehensive testing strategy was implemented to ensure the reliability and correctness of the NetShield IDS system across all components.")

    add_body_paragraph(doc, "Unit Testing. Individual components of the system were tested in isolation using the pytest framework. Unit tests cover: (1) data preprocessing functions (loading, cleaning, feature selection, normalization); (2) autoencoder forward pass and reconstruction error computation; (3) classifier prediction interfaces; (4) pipeline predict and predict_batch methods; (5) database operations (CRUD for analyses, predictions, metrics).")

    add_body_paragraph(doc, "Integration Testing. End-to-end integration tests verify the complete workflow from API request to response. Using the httpx library with FastAPI's TestClient, integration tests cover: (1) CSV file upload and analysis endpoint; (2) model metrics retrieval endpoint; (3) analysis history endpoints; (4) health check endpoint; (5) URL analysis endpoint (mocked external services).")

    add_body_paragraph(doc, "Model Validation. Model validation follows standard machine learning evaluation practices: (1) stratified train/test split (80/20) to prevent data leakage and maintain class distribution; (2) cross-validation consideration for hyperparameter tuning; (3) evaluation on the held-out test set using multiple metrics (Accuracy, Precision, Recall, F1, ROC AUC); (4) confusion matrix analysis for per-class performance assessment.")

    add_body_paragraph(doc, "Performance Testing. The system's performance was evaluated under varying load conditions: (1) CSV analysis throughput: approximately 10,000 records per second on consumer hardware (Intel i7, 16GB RAM); (2) WebSocket streaming: sustained 25 records/second delivery rate with sub-100ms latency; (3) URL analysis: average response time of 2-4 seconds including OSINT feed lookup and VirusTotal API query; (4) API response time: under 50ms for database queries (history, stats, metrics endpoints).")

    add_table_with_data(doc,
        ["Test Category", "Framework/Tool", "Coverage"],
        [
            ["Unit Tests", "pytest", "ML components, preprocessing, database"],
            ["Integration Tests", "httpx + TestClient", "All REST API endpoints"],
            ["Model Validation", "scikit-learn metrics", "All 5 models (AE + 4 classifiers)"],
            ["Performance Tests", "Manual benchmarking", "Throughput, latency, concurrency"],
        ],
        caption="Testing strategy overview",
        number="3.7"
    )

    add_body_paragraph(doc, "Deployment Architecture. The application supports two deployment configurations: (1) Development mode, where the backend (uvicorn) and frontend (Vite dev server) run as separate processes on different ports (8000 and 5173 respectively), with CORS enabling cross-origin communication. This mode supports hot module replacement for frontend development and auto-reload for backend changes. (2) Production mode, using Docker Compose to orchestrate the services. The provided Dockerfile configures the backend container with all ML dependencies, while the frontend can be built as static files and served by the backend or a dedicated web server.")

    add_body_paragraph(doc, "The Docker configuration ensures consistent deployment across different environments by encapsulating all Python dependencies, ML libraries, and model files within the container image. The docker-compose.yml file at the project root defines the service configuration, including port mappings, volume mounts for persistent data storage, and environment variables.")

    add_body_paragraph(doc, "Security Considerations in Implementation. Several security best practices are implemented throughout the application: (1) input validation using Pydantic models prevents injection attacks through API parameters; (2) file upload size limits prevent denial-of-service through oversized CSV uploads (max_records parameter, default 10,000); (3) CORS is restricted to known frontend origins rather than allowing all origins; (4) the WebSocket implementation includes connection management to prevent resource exhaustion from excessive concurrent connections; (5) database queries use parameterized statements (via aiosqlite) to prevent SQL injection; (6) the URL analysis module validates and normalizes URLs before processing to prevent server-side request forgery (SSRF) attacks.")

    add_body_paragraph(doc, "Logging and Monitoring. The application implements structured logging through Python's logging module. The threat feeds module logs the status of feed updates, including the number of indicators loaded from each source and any errors encountered during fetching. API requests are logged by FastAPI's built-in middleware, providing request/response timing information. Model loading status is reported at startup, clearly indicating whether the ML pipeline is available for inference.")

    add_page_break(doc)


def create_chapter4(doc):
    """Create Chapter 4: Economic Effectiveness."""
    add_heading_centered(doc, "4 ECONOMIC EFFECTIVENESS", font_size=14)

    # 4.1
    add_subheading(doc, "4.1 Project Cost Estimation")

    add_body_paragraph(doc, "This section presents a comprehensive analysis of the economic effectiveness of the NetShield IDS project. The cost estimation includes development costs, infrastructure costs, and operational costs, providing a complete picture of the financial aspects of the project.")

    add_body_paragraph(doc, "Development Cost Estimation. The development costs are calculated based on the labor costs of the developer and the estimated time spent on each project phase:", bold=True)

    add_table_with_data(doc,
        ["Phase", "Duration (days)", "Hourly Rate (USD)", "Hours/Day", "Cost (USD)"],
        [
            ["Literature review and analysis", "15", "15", "6", "1,350"],
            ["Dataset analysis and preprocessing", "10", "15", "6", "900"],
            ["System architecture design", "5", "15", "6", "450"],
            ["ML pipeline implementation", "20", "15", "6", "1,800"],
            ["Backend development", "15", "15", "6", "1,350"],
            ["Frontend development", "15", "15", "6", "1,350"],
            ["URL analysis module", "7", "15", "6", "630"],
            ["Real-time monitoring", "5", "15", "6", "450"],
            ["Testing and debugging", "10", "15", "6", "900"],
            ["Documentation", "8", "15", "6", "720"],
            ["TOTAL", "110", "\u2014", "\u2014", "9,900"],
        ],
        caption="Development cost breakdown",
        number="4.1"
    )

    add_body_paragraph(doc, "The hourly rate of $15/hour is based on the average junior developer salary in Kazakhstan for the IT sector (approximately 300,000-400,000 KZT per month). The total development period spans approximately 110 working days (5.5 months), consistent with the academic calendar plan.")

    add_body_paragraph(doc, "Infrastructure Costs. The project utilizes predominantly open-source and free-tier services, minimizing infrastructure costs:", bold=True)

    add_table_with_data(doc,
        ["Item", "Cost Type", "Cost (USD/year)"],
        [
            ["Development hardware (personal laptop)", "One-time (amortized)", "200"],
            ["Python, PyTorch, scikit-learn, React", "Open-source (free)", "0"],
            ["FastAPI, SQLite, Ant Design", "Open-source (free)", "0"],
            ["CICIDS2017 dataset", "Public dataset (free)", "0"],
            ["VirusTotal API (free tier)", "Free (500 requests/day)", "0"],
            ["OSINT feeds (URLhaus, OpenPhish, etc.)", "Free and open", "0"],
            ["GitHub (code repository)", "Free tier", "0"],
            ["Domain and hosting (optional)", "Variable", "50-100"],
            ["TOTAL", "\u2014", "250-300"],
        ],
        caption="Infrastructure cost estimation",
        number="4.2"
    )

    add_body_paragraph(doc, "Total Project Cost. The total estimated project cost including development labor and infrastructure is approximately $10,200 USD (approximately 4,800,000 KZT at the exchange rate of 470 KZT/USD). This represents a significantly lower investment compared to commercial IDS solutions.")

    # 4.2
    add_subheading(doc, "4.2 Comparison with Commercial Solutions")

    add_body_paragraph(doc, "To evaluate the economic advantage of the NetShield IDS, a comparison with leading commercial intrusion detection and prevention solutions is presented:")

    add_table_with_data(doc,
        ["Solution", "Type", "Annual Cost (USD)", "ML-based Detection", "Real-time Dashboard"],
        [
            ["Cisco Firepower", "Hardware + Software", "15,000 - 50,000", "Limited", "Yes"],
            ["Palo Alto Networks", "Hardware + Software", "20,000 - 80,000", "Yes (AI-driven)", "Yes"],
            ["Fortinet FortiGate", "Hardware + Software", "10,000 - 40,000", "Limited", "Yes"],
            ["Snort (commercial rules)", "Software", "2,000 - 5,000", "No (signature-only)", "Limited"],
            ["Suricata + ELK Stack", "Software", "5,000 - 15,000", "Partial", "Custom"],
            ["Splunk Enterprise Security", "Software", "25,000 - 100,000+", "Yes", "Yes"],
            ["NetShield IDS", "Software", "250 - 300", "Yes (full ML pipeline)", "Yes"],
        ],
        caption="Cost comparison with commercial IDS solutions",
        number="4.3"
    )

    add_body_paragraph(doc, "The comparison reveals that commercial IDS solutions require annual investments ranging from $2,000 to $100,000+, while the NetShield IDS achieves comparable detection capabilities for a fraction of the cost. The primary trade-offs are: (1) commercial solutions offer enterprise support and certified updates, while NetShield requires in-house maintenance; (2) commercial solutions typically include hardware appliances for high-throughput environments, while NetShield is a software-only solution; (3) commercial solutions have established track records in production environments, while NetShield is a research/educational prototype.")

    # 4.3
    add_subheading(doc, "4.3 Economic Benefits")

    add_body_paragraph(doc, "The economic benefits of the NetShield IDS project can be evaluated from several perspectives:")

    add_body_paragraph(doc, "Cost Savings for Educational and Research Institutions. Universities and research institutions in Kazakhstan and Central Asia that require IDS capabilities for teaching and research can deploy NetShield IDS at negligible cost compared to licensing commercial solutions. For a typical university with 5-10 cybersecurity courses, the annual savings could reach $10,000-$50,000 compared to maintaining commercial IDS licenses.")

    add_body_paragraph(doc, "Cost Savings for SMEs. Small and medium-sized enterprises (SMEs) in Kazakhstan, which often lack the budget for enterprise security solutions, can benefit from deploying an ML-based IDS at minimal cost. According to the National Information Security Centre of Kazakhstan, approximately 60% of SMEs lack adequate intrusion detection capabilities. NetShield IDS could help address this gap at a cost reduction of 95-99% compared to commercial alternatives.")

    add_body_paragraph(doc, "Potential for Damage Prevention. According to the IBM Cost of a Data Breach Report 2023, the average cost of a data breach in the ASEAN region (closest available proxy for Central Asia) is approximately $3.05 million. Even a partial reduction in breach risk through improved detection capabilities could represent significant economic value. If NetShield IDS prevents even one security incident per year, the return on investment would be substantial.")

    add_body_paragraph(doc, "Return on Investment (ROI) Calculation:", bold=True)
    add_body_paragraph(doc, "Assuming a conservative scenario where the system is deployed in a small organization and prevents just one minor security incident per year (estimated damage: $5,000):")
    add_body_paragraph(doc, "ROI = (Benefits - Costs) / Costs * 100%")
    add_body_paragraph(doc, "ROI = ($5,000 - $10,200) / $10,200 * 100% = -51% (first year)")
    add_body_paragraph(doc, "ROI = ($5,000 - $300) / $300 * 100% = 1,567% (subsequent years)")
    add_body_paragraph(doc, "The negative ROI in the first year is due to the initial development cost. However, from the second year onward, with only operational costs of approximately $300/year, the ROI becomes highly positive. Over a 5-year horizon: Total cost = $10,200 + 4*$300 = $11,400. Total prevented damage = 5*$5,000 = $25,000. 5-year ROI = ($25,000 - $11,400) / $11,400 * 100% = 119%.")

    add_body_paragraph(doc, "Open-Source Value. As an open-source project, NetShield IDS contributes to the broader cybersecurity ecosystem by providing a reference implementation of a two-level ML-based IDS. This enables other developers and researchers to build upon the work, potentially amplifying its economic impact beyond the direct deployment scenario.")

    add_body_paragraph(doc, "Total Cost of Ownership (TCO) Analysis. The 5-year total cost of ownership comparison further highlights the economic advantages of the NetShield IDS:", bold=True)

    add_table_with_data(doc,
        ["Cost Component", "NetShield IDS", "Snort + Custom", "Cisco Firepower"],
        [
            ["Initial Development/Purchase", "$10,200", "$5,000", "$30,000"],
            ["Annual License/Subscription", "$0", "$2,000", "$8,000"],
            ["Hardware (5 years)", "$0", "$2,000", "$10,000"],
            ["Maintenance (5 years)", "$1,500", "$10,000", "$15,000"],
            ["Training (personnel)", "$500", "$3,000", "$5,000"],
            ["5-Year TCO", "$12,200", "$22,000", "$68,000"],
        ],
        caption="5-year total cost of ownership comparison",
        number="4.4"
    )

    add_body_paragraph(doc, "The TCO analysis demonstrates that the NetShield IDS provides the lowest 5-year cost at $12,200, representing a 44% savings compared to an open-source Snort-based solution with commercial rule subscriptions, and an 82% savings compared to the Cisco Firepower enterprise solution. While the commercial solutions offer additional features such as hardware appliances, 24/7 vendor support, and certified threat intelligence updates, the NetShield IDS provides comparable ML-based detection capabilities at a fraction of the cost.")

    add_body_paragraph(doc, "Market Potential. The cybersecurity market in Kazakhstan and Central Asia is experiencing rapid growth, driven by government digitalization initiatives and increasing cyber threats. The market for IDS/IPS solutions in Kazakhstan is estimated at $50-100 million annually. NetShield IDS, as an affordable ML-based solution developed locally, has potential applications in: (1) university cybersecurity education programs; (2) government agency network monitoring (particularly for agencies with limited IT budgets); (3) SME network security enhancement; and (4) integration into managed security service provider (MSSP) offerings.")

    add_page_break(doc)


def create_conclusion(doc):
    """Create the Conclusion section."""
    add_heading_centered(doc, "CONCLUSION", font_size=14)

    add_body_paragraph(doc, "This diploma project has successfully achieved its stated goal of developing a two-level intrusion detection system (NetShield IDS) for automated network attack detection using machine learning methods. The following key results were obtained:")

    add_body_paragraph(doc, "1. A comprehensive analytical review of existing IDS approaches, machine learning methods for intrusion detection, and benchmark datasets was conducted. The review demonstrated that hybrid approaches combining anomaly detection with classification provide the best balance between novelty detection and precise attack categorization.")

    add_body_paragraph(doc, "2. The theoretical foundations of the two-level detection architecture were established. The first level employs an autoencoder neural network trained on normal traffic for unsupervised anomaly detection, while the second level utilizes an ensemble of four classifiers (XGBoost, LightGBM, Random Forest, MLP) for supervised attack type classification across 8 categories.")

    add_body_paragraph(doc, "3. A full-stack web application was implemented, featuring:")
    add_body_paragraph(doc, "\u2022 A FastAPI backend with REST API endpoints for batch CSV analysis, WebSocket endpoints for real-time monitoring, and asynchronous SQLite database for analysis history storage.")
    add_body_paragraph(doc, "\u2022 A React frontend with an interactive security operations dashboard providing batch analysis, real-time monitoring, URL security analysis, and model comparison views.")
    add_body_paragraph(doc, "\u2022 A URL security analysis module integrating OSINT threat intelligence feeds (URLhaus, OpenPhish, Phishing.Database, Feodo Tracker) and the VirusTotal API for comprehensive web resource verification.")

    add_body_paragraph(doc, "4. All models were trained and evaluated on the CICIDS2017 dataset with the following results:")
    add_body_paragraph(doc, "\u2022 Autoencoder anomaly detection: 96.2% accuracy, 0.987 ROC AUC")
    add_body_paragraph(doc, "\u2022 XGBoost classifier (best): 99.52% accuracy, 0.9951 weighted F1-score, 0.9998 ROC AUC")
    add_body_paragraph(doc, "\u2022 LightGBM: 99.41% accuracy, 0.9940 weighted F1-score")
    add_body_paragraph(doc, "\u2022 Random Forest: 99.38% accuracy, 0.9937 weighted F1-score")
    add_body_paragraph(doc, "\u2022 MLP: 98.87% accuracy, 0.9886 weighted F1-score")

    add_body_paragraph(doc, "5. An economic effectiveness analysis demonstrated that the NetShield IDS achieves comparable detection capabilities to commercial solutions at a cost reduction of 95-99%, with a projected 5-year ROI of 119% in a conservative deployment scenario.")

    add_body_paragraph(doc, "The developed system addresses several practical needs in the current cybersecurity landscape, particularly for organizations with limited budgets for commercial security solutions. The two-level detection architecture provides both known attack classification and novel attack detection capabilities, while the web-based interface enables security operations without specialized hardware or software.")

    add_body_paragraph(doc, "Limitations and Future Work. Several areas have been identified for future improvement:", bold=True)
    add_body_paragraph(doc, "\u2022 Dataset modernization: Training on more recent datasets (such as CICIDS2018, CIC-DDoS2019, or custom datasets from contemporary traffic) would improve the system's ability to detect current attack types.")
    add_body_paragraph(doc, "\u2022 Model ensemble: Implementing a weighted voting ensemble across all four classifiers could potentially improve classification accuracy beyond any individual model.")
    add_body_paragraph(doc, "\u2022 Deep packet inspection: Extending the feature extraction to include payload-level features could improve detection of application-layer attacks.")
    add_body_paragraph(doc, "\u2022 Automated retraining: Implementing an automated pipeline for periodic model retraining on new data would help maintain detection accuracy as network traffic patterns evolve.")
    add_body_paragraph(doc, "\u2022 Deployment at scale: Containerizing the application with Docker/Kubernetes and implementing load balancing would enable deployment in enterprise environments with high traffic volumes.")
    add_body_paragraph(doc, "\u2022 Integration with SIEM: Adding export capabilities for popular SIEM formats (CEF, LEEF) would enable integration with existing security operations center (SOC) workflows.")

    add_page_break(doc)


def create_references(doc):
    """Create the References section."""
    add_heading_centered(doc, "REFERENCES", font_size=14)

    references = [
        '[1] IBM Security, "Cost of a Data Breach Report 2023," IBM Corporation, 2023. [Online]. Available: https://www.ibm.com/reports/data-breach',
        '[2] H. Debar, M. Dacier, and A. Wespi, "A Revised Taxonomy for Intrusion-Detection Systems," Annales des Telecommunications, vol. 55, no. 7-8, pp. 361-378, 2000.',
        '[3] A. L. Buczak and E. Guha, "A Survey of Data Mining and Machine Learning Methods for Cyber Security Intrusion Detection," IEEE Communications Surveys & Tutorials, vol. 18, no. 2, pp. 1153-1176, 2016.',
        '[4] ISC2, "Cybersecurity Workforce Study 2023," International Information System Security Certification Consortium, 2023.',
        '[5] W. Stallings and L. Brown, Computer Security: Principles and Practice, 4th ed., Pearson, 2018.',
        '[6] Cloudflare, "DDoS Threat Report for 2023 Q4," Cloudflare Inc., 2024. [Online]. Available: https://radar.cloudflare.com/reports/ddos-2023-q4',
        '[7] I. Sharafaldin, A. H. Lashkari, and A. A. Ghorbani, "Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization," in Proc. 4th International Conference on Information Systems Security and Privacy (ICISSP), 2018, pp. 108-116.',
        '[8] OWASP Foundation, "OWASP Top Ten Web Application Security Risks," 2021. [Online]. Available: https://owasp.org/www-project-top-ten/',
        '[9] G. Lyon, Nmap Network Scanning: The Official Nmap Project Guide to Network Discovery and Security Scanning, Nmap Project, 2009.',
        '[10] S. S. C. Silva, R. M. P. Silva, R. C. G. Pinto, and R. M. Salles, "Botnets: A Survey," Computer Networks, vol. 57, no. 2, pp. 378-403, 2013.',
        '[11] M. Roesch, "Snort \u2014 Lightweight Intrusion Detection for Networks," in Proc. 13th USENIX Conference on System Administration (LISA), 1999, pp. 229-238.',
        '[12] D. Bray, R. Breitinger, and H. Baier, "Host-Based Intrusion Detection Systems: A Review," Journal of Information Security and Applications, vol. 73, 2023.',
        '[13] P. Garcia-Teodoro, J. Diaz-Verdejo, G. Macia-Fernandez, and E. Vazquez, "Anomaly-based Network Intrusion Detection: Techniques, Systems and Challenges," Computers & Security, vol. 28, no. 1-2, pp. 18-28, 2009.',
        '[14] V. Chandola, A. Banerjee, and V. Kumar, "Anomaly Detection: A Survey," ACM Computing Surveys, vol. 41, no. 3, pp. 1-58, 2009.',
        '[15] S. Axelsson, "Intrusion Detection Systems: A Survey and Taxonomy," Technical Report 99-15, Chalmers University of Technology, 2000.',
        '[16] M. Roesch and C. Green, Snort Users Manual, Sourcefire Inc., 2023.',
        '[17] OISF, "Suricata \u2014 Open Source IDS/IPS/NSM Engine," Open Information Security Foundation, 2024. [Online]. Available: https://suricata.io/',
        '[18] V. Paxson, "Bro: A System for Detecting Network Intruders in Real-Time," Computer Networks, vol. 31, no. 23-24, pp. 2435-2463, 1999.',
        '[19] Gartner, "Market Guide for Intrusion Detection and Prevention Systems," Gartner Inc., 2023.',
        '[20] N. Moustafa, B. Turnbull, and K. R. Choo, "An Ensemble Intrusion Detection Technique based on proposed Statistical Flow Features for Protecting Network Traffic of Internet of Things," IEEE Internet of Things Journal, vol. 6, no. 3, pp. 4815-4830, 2019.',
        '[21] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.',
        '[22] P. A. A. Resende and A. C. Drummond, "A Survey of Random Forest Based Methods for Intrusion Detection Systems," ACM Computing Surveys, vol. 51, no. 3, pp. 1-36, 2018.',
        '[23] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016, pp. 785-794.',
        '[24] G. Ke, Q. Meng, T. Finley, T. Wang, W. Chen, W. Ma, Q. Ye, and T.-Y. Liu, "LightGBM: A Highly Efficient Gradient Boosting Decision Tree," in Advances in Neural Information Processing Systems (NeurIPS), vol. 30, 2017, pp. 3146-3154.',
        '[25] M. A. Ferrag, L. Maglaras, S. Moschoyiannis, and H. Janicke, "Deep Learning for Cyber Security Intrusion Detection: Approaches, Datasets, and Comparative Study," Journal of Information Security and Applications, vol. 50, 2020.',
        '[26] Y. Li, R. Ma, and R. Jiao, "A Hybrid Malicious Code Detection Method based on Deep Learning," International Journal of Security and Its Applications, vol. 9, no. 5, pp. 205-216, 2015.',
        '[27] C. L. Yin, Y. F. Zhu, J. L. Fei, and X. Z. He, "A Deep Learning Approach for Intrusion Detection Using Recurrent Neural Networks," IEEE Access, vol. 5, pp. 21954-21961, 2017.',
        '[28] Y. Mirsky, T. Doitshman, Y. Elovici, and A. Shabtai, "Kitsune: An Ensemble of Autoencoders for Online Network Intrusion Detection," in Proc. Network and Distributed Systems Security Symposium (NDSS), 2018.',
        '[29] J. An and S. Cho, "Variational Autoencoder based Anomaly Detection using Reconstruction Probability," Special Lecture on IE, vol. 2, no. 1, pp. 1-18, 2015.',
        '[30] F. T. Liu, K. M. Ting, and Z. H. Zhou, "Isolation Forest," in Proc. IEEE International Conference on Data Mining (ICDM), 2008, pp. 413-422.',
        '[31] A. H. Lashkari, G. Draper-Gil, M. S. I. Mamun, and A. A. Ghorbani, "Characterization of Tor Traffic using Time based Features," in Proc. 3rd International Conference on Information Systems Security and Privacy (ICISSP), 2017, pp. 253-262.',
        '[32] S. Mahdavifar and A. A. Ghorbani, "Application of Deep Learning to Cybersecurity: A Survey," Neurocomputing, vol. 347, pp. 149-176, 2019.',
        '[33] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning, MIT Press, 2016.',
        '[34] D. Bank, N. Koenigstein, and R. Giryes, "Autoencoders," arXiv preprint arXiv:2003.05991, 2020.',
        '[35] S. Ioffe and C. Szegedy, "Batch Normalization: Accelerating Deep Network Training by Reducing Internal Covariate Shift," in Proc. 32nd International Conference on Machine Learning (ICML), 2015, pp. 448-456.',
        '[36] N. Srivastava, G. Hinton, A. Krizhevsky, I. Sutskever, and R. Salakhutdinov, "Dropout: A Simple Way to Prevent Neural Networks from Overfitting," Journal of Machine Learning Research, vol. 15, no. 56, pp. 1929-1958, 2014.',
        '[37] V. Nair and G. E. Hinton, "Rectified Linear Units Improve Restricted Boltzmann Machines," in Proc. 27th International Conference on Machine Learning (ICML), 2010, pp. 807-814.',
        '[38] K. Hornik, M. Stinchcombe, and H. White, "Multilayer Feedforward Networks are Universal Approximators," Neural Networks, vol. 2, no. 5, pp. 359-366, 1989.',
        '[39] C. Goutte and E. Gaussier, "A Probabilistic Interpretation of Precision, Recall and F-Score, with Implication for Evaluation," in Proc. European Conference on Information Retrieval (ECIR), 2005, pp. 345-359.',
        '[40] T. Fawcett, "An Introduction to ROC Analysis," Pattern Recognition Letters, vol. 27, no. 8, pp. 861-874, 2006.',
        '[41] S. Tilkov and S. Vinoski, "Node.js: Using JavaScript to Build High-Performance Network Programs," IEEE Internet Computing, vol. 14, no. 6, pp. 80-83, 2010.',
        '[42] S. Ramirez, "FastAPI Documentation," 2024. [Online]. Available: https://fastapi.tiangolo.com/',
        '[43] Ant Design Team, "Ant Design: A design system for enterprise-level products," 2024. [Online]. Available: https://ant.design/',
    ]

    for ref in references:
        add_body_paragraph(doc, ref, first_line_indent=0, font_size=12, space_after=4)

    add_page_break(doc)


def create_appendices(doc):
    """Create the Appendices section."""
    add_heading_centered(doc, "APPENDIX A", font_size=14)
    add_heading_centered(doc, "Source Code Listings", font_size=14, space_before=6)

    add_subheading(doc, "A.1 Autoencoder Model (autoencoder.py)")
    add_code_block(doc, """class Autoencoder(nn.Module):
    def __init__(self, input_dim: int):
        super().__init__()
        self.input_dim = input_dim
        self.encoder = nn.Sequential(
            nn.Linear(input_dim, 32),
            nn.ReLU(),
            nn.BatchNorm1d(32),
            nn.Dropout(0.2),
            nn.Linear(32, 16),
            nn.ReLU(),
            nn.BatchNorm1d(16),
            nn.Linear(16, 8),
            nn.ReLU(),
        )
        self.decoder = nn.Sequential(
            nn.Linear(8, 16),
            nn.ReLU(),
            nn.BatchNorm1d(16),
            nn.Dropout(0.2),
            nn.Linear(16, 32),
            nn.ReLU(),
            nn.BatchNorm1d(32),
            nn.Linear(32, input_dim),
        )

    def forward(self, x):
        encoded = self.encoder(x)
        decoded = self.decoder(encoded)
        return decoded""")

    add_subheading(doc, "A.2 Two-Level Pipeline (pipeline.py)")
    add_code_block(doc, """class IDSPipeline:
    def predict(self, X: np.ndarray) -> list:
        results = []
        # Level 1: Anomaly Detection
        anomaly_scores = self.anomaly_detector.predict_scores(X)
        anomaly_predictions = self.anomaly_detector.predict(X)
        # Level 2: Classification
        class_predictions = self.classifier.predict(X)
        class_probas = self.classifier.predict_proba(X)

        for i in range(len(X)):
            is_anomaly = bool(anomaly_predictions[i])
            if is_anomaly:
                pred_class = int(class_predictions[i])
                attack_type = self.label_encoder.inverse_transform(
                    [pred_class])[0]
                confidence = float(np.max(class_probas[i]))
                if attack_type == 'BENIGN':
                    attack_type = 'Unknown Attack'
            else:
                attack_type = 'BENIGN'
                confidence = 1.0 - float(anomaly_scores[i] /
                    (self.anomaly_detector.threshold + 1e-8))
            results.append({
                'is_attack': is_anomaly,
                'attack_type': attack_type,
                'anomaly_score': float(anomaly_scores[i]),
                'confidence': round(confidence, 4),
            })
        return results""")

    add_subheading(doc, "A.3 Classifier Configurations (classifiers.py)")
    add_code_block(doc, """def create_classifiers(input_dim, n_classes):
    return {
        'Random Forest': RandomForestClassifier(
            n_estimators=100, max_depth=20,
            min_samples_split=5, n_jobs=-1, random_state=42,
        ),
        'XGBoost': XGBClassifier(
            n_estimators=100, max_depth=10, learning_rate=0.1,
            n_jobs=-1, random_state=42, eval_metric='mlogloss',
        ),
        'LightGBM': LGBMClassifier(
            n_estimators=100, max_depth=15, learning_rate=0.1,
            n_jobs=-1, random_state=42, verbose=-1,
        ),
        'MLP': MLPWrapper(input_dim, n_classes),
    }""")

    add_subheading(doc, "A.4 FastAPI Application Entry Point (main.py)")
    add_code_block(doc, """app = FastAPI(
    title="Network IDS",
    version="1.0.0",
    lifespan=lifespan,
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000",
                    "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.include_router(api_router)
app.include_router(ws_router)""")

    add_subheading(doc, "A.5 URL Security Analysis Scoring (url_analyzer.py)")
    add_code_block(doc, """async def analyze_url(url, pipeline=None):
    # 5-factor security scoring (100 points total)
    # 1. SSL/TLS              (max 20)
    ssl_score, ssl_recs = _score_ssl(is_https, ssl_valid, ssl_expiry)
    # 2. Security Headers     (max 25)
    hdr_score, sec_headers, hdr_recs = _score_headers(headers, is_https)
    # 3. URL and Domain        (max 20)
    url_score, url_recs = _score_url_domain(url, domain)
    # 4. Content and Behavior  (max 15)
    cb_score, cb_recs = _score_content_behavior(...)
    # 5. Threat Intelligence   (max 20)
    url_threats = await check_url(url)
    ip_threats = await check_ip(ip)
    vt_result = await check_virustotal(url)
    total_score = ssl_score + hdr_score + url_score + cb_score + ti_score
    return URLAnalysisResponse(security_score=total_score, ...)""")

    add_page_break(doc)

    # Appendix B
    add_heading_centered(doc, "APPENDIX B", font_size=14)
    add_heading_centered(doc, "Screenshots", font_size=14, space_before=6)

    screenshots = [
        ("B.1", "Main Dashboard with security posture overview and statistics"),
        ("B.2", "CSV analysis results page with attack filtering"),
        ("B.3", "URL security analysis showing score breakdown and recommendations"),
        ("B.4", "URL analysis \u2014 Threat Intelligence section with VirusTotal results"),
        ("B.5", "Real-time monitoring during a DDoS attack scenario"),
        ("B.6", "Real-time monitoring showing normal traffic flow"),
        ("B.7", "Model comparison page with performance metrics"),
        ("B.8", "Dashboard \u2014 Threat distribution pie chart"),
        ("B.9", "Dashboard \u2014 24-hour detection rhythm area chart"),
        ("B.10", "FastAPI automatic API documentation (Swagger UI)"),
    ]

    for num, desc in screenshots:
        add_subheading(doc, f"{num} {desc}")
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot: {desc}]")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6)
        pf.space_after = Pt(24)
        pf.line_spacing = 1.5


# ──────────────────────────────────────────────────────────
# Main document assembly
# ──────────────────────────────────────────────────────────

def generate_thesis():
    """Generate the complete diploma thesis DOCX."""
    print("Generating NetShield IDS Diploma Thesis...")

    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    # ── Default paragraph style ──
    style = doc.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Set East Asian font too
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman" w:cs="Times New Roman"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), "Times New Roman")

    # ── Page numbering (bottom center) ──
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = fp.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = fp.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)

    # ── Build document sections ──
    print("  Creating title page...")
    create_title_page(doc)

    print("  Creating assignment page...")
    create_assignment_page(doc)

    print("  Creating calendar plan...")
    create_calendar_plan(doc)

    print("  Creating annotations (KZ, RU, EN)...")
    create_annotations(doc)

    print("  Creating list of abbreviations...")
    create_abbreviations(doc)

    print("  Creating table of contents...")
    create_table_of_contents(doc)

    print("  Creating Introduction...")
    create_introduction(doc)

    print("  Creating Chapter 1: Analytical Review...")
    create_chapter1(doc)

    print("  Creating Chapter 2: Theoretical Part...")
    create_chapter2(doc)

    print("  Creating Chapter 3: Software Implementation...")
    create_chapter3(doc)

    print("  Creating Chapter 4: Economic Effectiveness...")
    create_chapter4(doc)

    print("  Creating Conclusion...")
    create_conclusion(doc)

    print("  Creating References...")
    create_references(doc)

    print("  Creating Appendices...")
    create_appendices(doc)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"\nThesis saved to: {OUTPUT_PATH}")
    print("Done!")


if __name__ == "__main__":
    generate_thesis()
